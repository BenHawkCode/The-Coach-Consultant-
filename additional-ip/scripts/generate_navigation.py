"""
Generate Navigation Documents for Claude Projects Client Pack.
Creates a master index at the root level plus a folder index inside each subfolder
so clients can easily navigate between documents in Google Drive.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "Claude-Projects-Client-Pack")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. Unauthorised sharing, copying, reproduction, or sale is prohibited without written permission. Legal action will be taken for infringements. Personal use is allowed for active members or opt-ins only."
FONT_NAME = "Poppins"


# ─── Complete pack contents ───

FOLDERS = [
    {
        "folder": "00-Onboarding",
        "title": "Onboarding",
        "description": "Everything you need to get started, set up your account, and understand the system.",
        "next": "Next: Open the project folder that matches your biggest need. Check 09-ALL-34-PROJECTS-AT-A-GLANCE first to decide.",
        "reading_order": True,
        "files": [
            ("00-COST-SAVINGS-BREAKDOWN.docx", "What this system replaces: \u00a362,450/month in staff wages broken down by role"),
            ("01-YOUR-FIRST-15-MINUTES.docx", "Set up your first project in 15 minutes with zero knowledge files needed"),
            ("02-START-HERE.docx", "Full orientation guide with screenshots showing how to set up any Claude Project"),
            ("03-DEFINE-YOUR-BRAND-FIRST.docx", "Fill-in worksheet for the 7 placeholders used in every project template"),
            ("04-WHY-YOU-NEED-PRO-AND-MAX.docx", "Which Claude plan to choose and why the free plan will not work"),
            ("05-RECOMMENDED-TOOLS-AND-SETUP.docx", "Wispr Flow, browser setup, and how to prepare your knowledge files"),
            ("06-KNOWLEDGE-FILE-STARTER-KIT.docx", "How to find, create, and organise the files your projects need"),
            ("07-YOUR-FIRST-WEEK-FRAMEWORK.docx", "Day-by-day setup plan plus Month 1 and Quarter 1 rollout sequence"),
            ("08-HOW-TO-GET-BEST-RESULTS.docx", "5 rules for writing prompts that get you what you want first time"),
            ("09-ALL-34-PROJECTS-AT-A-GLANCE.docx", "Every project with ROI ratings, difficulty levels, and Top 5 recommendations"),
            ("10-SETUP-PROGRESS-TRACKER.docx", "Tick-off checklist to track your setup across 5 phases"),
            ("11-TROUBLESHOOTING-AND-FAQS.docx", "Quick fixes for the most common issues"),
            ("12-PROJECT-REVIEW-CHECKLIST.docx", "Per-project setup checklist for all 34 projects"),
        ],
    },
    {
        "folder": "01-Content-Marketing",
        "title": "Content Marketing",
        "description": "Projects for planning, writing, and managing content across your business.",
        "next": "Next: 02-Social-Media for platform-specific content, or back to 00-Onboarding if you need more setup help.",
        "files": [
            ("01-Content-Strategy-Calendar.docx", "Project 1: Plan monthly themes and weekly content calendars"),
            ("02-Blog-Long-Form-Content.docx", "Project 2: SEO articles, guides, and thought leadership pieces"),
            ("03-Email-Marketing-Sequences.docx", "Project 3: Welcome sequences, nurture campaigns, launch emails"),
            ("04-Lead-Magnets-Opt-In.docx", "Project 4: Checklists, guides, and downloadable resources"),
            ("33-Community-Management.docx", "Project 33: Group rules, welcome sequences, engagement content"),
        ],
    },
    {
        "folder": "02-Social-Media",
        "title": "Social Media",
        "description": "Projects for creating content on Instagram, LinkedIn, YouTube, and podcasts.",
        "next": "Next: 03-Sales-Lead-Generation to turn your audience into clients.",
        "files": [
            ("05-Instagram-Content-Creator.docx", "Project 5: Captions, carousels, reel scripts, stories"),
            ("06-LinkedIn-Thought-Leadership.docx", "Project 6: Posts, articles, connection messages"),
            ("07-YouTube-Content-Scripts.docx", "Project 7: Full scripts with timestamps, titles, descriptions"),
            ("08-Content-Repurposing-Engine.docx", "Project 8: Turn one piece of content into five platforms"),
            ("32-Podcast-Strategy-Guest-Appearances.docx", "Project 32: Topic lists, speaker bios, podcast pitches"),
        ],
    },
    {
        "folder": "03-Sales-Lead-Generation",
        "title": "Sales & Lead Generation",
        "description": "Projects for writing sales copy, proposals, outreach scripts, and webinar content.",
        "next": "Next: 04-Client-Fulfilment to deliver a great experience once they sign up.",
        "files": [
            ("09-Sales-Landing-Page-Copy.docx", "Project 9: High-converting sales pages and landing pages"),
            ("10-Proposal-Pitch-Writer.docx", "Project 10: Client proposals, pitch decks, follow-ups"),
            ("11-DM-Outreach-Sequences.docx", "Project 11: Personalised outreach and follow-up scripts"),
            ("12-Webinar-Workshop-Planner.docx", "Project 12: Webinar structure, scripts, follow-up sequences"),
        ],
    },
    {
        "folder": "04-Client-Fulfilment",
        "title": "Client Fulfilment & Delivery",
        "description": "Projects for onboarding, session prep, reporting, course creation, and retention.",
        "next": "Next: 05-Operations to streamline how your business runs behind the scenes.",
        "files": [
            ("13-Client-Onboarding-Assistant.docx", "Project 13: Welcome packets, onboarding emails, intake forms"),
            ("14-Coaching-Session-Prep-Notes.docx", "Project 14: Session agendas, questions, follow-up summaries"),
            ("15-Client-Reporting-Results.docx", "Project 15: Monthly reports, progress summaries, ROI breakdowns"),
            ("16-Course-Programme-Content.docx", "Project 16: Module outlines, lesson scripts, workbook pages"),
            ("31-Client-Offboarding-Retention.docx", "Project 31: Offboarding sequences, referral systems, alumni offers"),
        ],
    },
    {
        "folder": "05-Operations",
        "title": "Operations & Business Management",
        "description": "Projects for SOPs, reporting, team management, finance, and automation.",
        "next": "Next: 06-Strategy-Growth for higher-level business planning and positioning.",
        "files": [
            ("17-SOP-Process-Documentation.docx", "Project 17: Step-by-step procedures and training guides"),
            ("18-KPI-Business-Reporting.docx", "Project 18: Metrics analysis, performance reports, trend spotting"),
            ("19-Team-Communication-Management.docx", "Project 19: Team updates, meeting agendas, feedback reviews"),
            ("20-Finance-Invoicing-Assistant.docx", "Project 20: Pricing structures, invoice templates, budgets"),
            ("29-Finance-Strategic-Reporting.docx", "Project 29: Revenue breakdowns, forecasting, CFO-level analysis"),
            ("34-Tech-Stack-Automation-Setup.docx", "Project 34: Tool recommendations, automation workflows"),
        ],
    },
    {
        "folder": "06-Strategy-Growth",
        "title": "Strategy & Growth",
        "description": "Projects for business planning, brand positioning, partnerships, research, and offer design.",
        "next": "Next: 07-Admin-Organisation for email, meetings, legal, and personal productivity.",
        "files": [
            ("21-Business-Strategy-Planning.docx", "Project 21: Quarterly plans, goal setting, business vision"),
            ("22-Brand-Positioning-Strategist.docx", "Project 22: Brand identity, positioning, differentiation"),
            ("23-Partnership-Collaboration.docx", "Project 23: JV partners, cross-promotion, collaboration planning"),
            ("24-Customer-Research-Insights.docx", "Project 24: Market research, customer insights, competitive analysis"),
            ("30-Offer-Design-Pricing-Strategy.docx", "Project 30: Offer architecture, pricing, offer waterfall"),
        ],
    },
    {
        "folder": "07-Admin-Organisation",
        "title": "Admin & Organisation",
        "description": "Projects for email management, meetings, legal documents, and personal productivity.",
        "next": "Next: 08-Claude-Code-Guides for managing and bulk-updating all your projects at scale.",
        "files": [
            ("25-Email-Communication-Manager.docx", "Project 25: Email templates, FAQs, communication standards"),
            ("26-Meeting-Event-Coordinator.docx", "Project 26: Agendas, event checklists, follow-up templates"),
            ("27-Legal-Compliance-Drafts.docx", "Project 27: Terms, privacy policies, disclaimers, contracts"),
            ("28-Personal-Productivity-Planning.docx", "Project 28: Weekly planning, priorities, goal setting"),
        ],
    },
    {
        "folder": "08-Claude-Code-Guides",
        "title": "Claude Code Guides",
        "description": "Ready-to-use prompts for managing, updating, and maintaining all your Claude Projects at scale using Claude Code.",
        "next": "You have explored the full pack. Go back to 00-Onboarding/10-SETUP-PROGRESS-TRACKER to track your progress.",
        "files": [
            ("00-WHAT-IS-CLAUDE-CODE.docx", "What Claude Code is, when to use it, and how to use the prompts in this folder"),
            ("01-FULL-REBRAND-UPDATE.docx", "Update brand files, names, and messaging across all projects at once"),
            ("02-BULK-CTA-UPDATE.docx", "Change your website URL, booking link, or CTA across every project"),
            ("03-PRICING-AND-OFFER-UPDATE.docx", "Update pricing, launch new offers, retire old ones across all projects"),
            ("04-TARGET-AUDIENCE-UPDATE.docx", "Push audience or niche changes to every project"),
            ("05-QUARTERLY-HEALTH-CHECK.docx", "Full audit of all projects for accuracy, consistency, and optimisation"),
            ("06-NEW-KNOWLEDGE-FILE-ROLLOUT.docx", "Add a new document to the right projects without missing any"),
            ("07-VOICE-AND-TONE-RECALIBRATION.docx", "Recalibrate your voice across all writing projects"),
            ("08-SEASONAL-CAMPAIGN-SETUP.docx", "Temporarily update projects for a launch or promo with revert guide"),
            ("09-PROJECT-BACKUP-AND-EXPORT.docx", "Create a full snapshot of all project configurations"),
            ("10-NEW-PROJECT-SETUP-GUIDED.docx", "Claude Code walks you through setting up a new project step by step"),
            ("11-COMPLIANCE-AND-DISCLAIMER-UPDATE.docx", "Audit and update legal disclaimers across all projects"),
            ("12-TEAM-VA-HANDOVER-PACK.docx", "Generate a complete handover document for a team member or VA"),
        ],
    },
]


# ─── Styling helpers ───

def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for style_name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = FONT_NAME


def new_doc():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)
    return doc


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_file_entry(doc, filename, description):
    """Add a single file entry: bold blue filename + grey description."""
    p = doc.add_paragraph()

    run_name = p.add_run(filename)
    run_name.font.name = FONT_NAME
    run_name.font.size = Pt(10)
    run_name.bold = True
    run_name.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)

    run_sep = p.add_run("  \u2014  ")
    run_sep.font.name = FONT_NAME
    run_sep.font.size = Pt(10)
    run_sep.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    run_desc = p.add_run(description)
    run_desc.font.name = FONT_NAME
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_nav_hint(doc, text):
    """Add a navigation hint in green bold."""
    p = doc.add_paragraph()
    run = p.add_run(f"\u2192  {text}")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)


# ─── Master Index ───

def create_master_index():
    doc = new_doc()

    add_heading(doc, "Claude Projects Client Pack")
    add_para(doc, "Master Index", bold=True, size=16)
    doc.add_paragraph("")

    add_para(doc, "This is your map to everything in this pack. Use it to find any document quickly.", size=12)
    doc.add_paragraph("")

    add_heading(doc, "Where to Start", level=2)
    add_para(doc, "Open the 00-Onboarding folder and read the documents in order (00 through 12). This will get you set up, give you a quick win, and show you how to use the full system.", bold=True)
    doc.add_paragraph("")

    add_para(doc, "Reading order for onboarding:", bold=True)
    add_para(doc, "00 Cost Savings \u2192 01 First 15 Minutes \u2192 02 Start Here \u2192 03 Define Your Brand \u2192 04 Pro & Max Guide \u2192 05 Tools Setup \u2192 06 Knowledge Files \u2192 07 First Week Plan \u2192 08 Prompting Guide \u2192 09 All Projects Overview \u2192 10 Progress Tracker \u2192 11 Troubleshooting \u2192 12 Review Checklist", size=10, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Each folder section
    for folder_data in FOLDERS:
        doc.add_page_break()
        add_heading(doc, f"{folder_data['folder']}: {folder_data['title']}", level=2)
        add_para(doc, folder_data['description'], size=11, color=RGBColor(0x55, 0x55, 0x55))
        doc.add_paragraph("")

        for filename, desc in folder_data['files']:
            add_file_entry(doc, filename, desc)

        doc.add_paragraph("")
        add_nav_hint(doc, folder_data['next'])

    # Summary
    doc.add_page_break()
    add_heading(doc, "Pack Summary", level=2)
    doc.add_paragraph("")

    total_files = sum(len(f['files']) for f in FOLDERS)
    add_para(doc, f"Total documents: {total_files}", bold=True)
    add_para(doc, "9 folders covering onboarding, content, social media, sales, client fulfilment, operations, strategy, admin, and Claude Code guides.")
    doc.add_paragraph("")
    add_para(doc, "34 ready-to-use Claude Project templates with copy-paste instructions, knowledge file checklists, and ROI ratings.", bold=True)
    doc.add_paragraph("")
    add_para(doc, "Questions? Check 00-Onboarding/11-TROUBLESHOOTING-AND-FAQS for quick fixes to common issues.")

    # Save to root
    os.makedirs(BASE_DIR, exist_ok=True)
    filepath = os.path.join(BASE_DIR, "00-MASTER-INDEX.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ─── Folder Indexes ───

def create_folder_index(folder_data):
    doc = new_doc()

    add_heading(doc, f"{folder_data['title']}")
    add_para(doc, "Folder Contents", bold=True, size=14)
    doc.add_paragraph("")

    add_nav_hint(doc, "Back to Master Index: Open 00-MASTER-INDEX.docx in the main folder")
    doc.add_paragraph("")

    add_para(doc, folder_data['description'], size=11, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Reading order for onboarding
    if folder_data.get('reading_order'):
        add_heading(doc, "Recommended Reading Order", level=2)
        add_para(doc, "Read these documents in order for the best experience:", bold=True)
        doc.add_paragraph("")
        for i, (filename, desc) in enumerate(folder_data['files']):
            p = doc.add_paragraph()
            run_num = p.add_run(f"{i + 1}.  ")
            run_num.font.name = FONT_NAME
            run_num.font.size = Pt(10)
            run_num.bold = True
            run_num.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)

            run_name = p.add_run(filename)
            run_name.font.name = FONT_NAME
            run_name.font.size = Pt(10)
            run_name.bold = True

            run_desc = p.add_run(f"  \u2014  {desc}")
            run_desc.font.name = FONT_NAME
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        add_heading(doc, "Documents in This Folder", level=2)
        doc.add_paragraph("")
        for filename, desc in folder_data['files']:
            add_file_entry(doc, filename, desc)

    doc.add_paragraph("")
    doc.add_paragraph("")
    add_heading(doc, "Where to Go Next", level=2)
    add_nav_hint(doc, folder_data['next'])

    # Save into folder
    folder_path = os.path.join(BASE_DIR, folder_data['folder'])
    os.makedirs(folder_path, exist_ok=True)
    filepath = os.path.join(folder_path, "00-FOLDER-INDEX.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ─── Generate everything ───

def main():
    print("Generating Navigation Documents...\n")

    create_master_index()
    for folder_data in FOLDERS:
        create_folder_index(folder_data)

    print(f"\nDone! 10 navigation documents created (1 master + 9 folder indexes)")


if __name__ == "__main__":
    main()
