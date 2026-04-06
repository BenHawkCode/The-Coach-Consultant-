"""
Generate Cost Savings Breakdown .docx with branded styling.
Uses same logo, Poppins font, header, and footer as all other client pack documents.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "Claude-Projects-Client-Pack")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. Unauthorised sharing, copying, reproduction, or sale is prohibited without written permission. Legal action will be taken for infringements. Personal use is allowed for active members or opt-ins only."
FONT_NAME = "Poppins"


def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for style_name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = FONT_NAME


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = FONT_NAME
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    return table


def main():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)

    # ── TITLE ──
    add_heading(doc, "What This System Replaces")
    add_para(doc, "The Real Cost of Hiring a Team vs Using Claude Projects", size=14, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # ── BIG NUMBER ──
    add_para(doc, "TOTAL MONTHLY STAFF COST YOU'RE REPLACING:", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "\u00a362,450/month", bold=True, size=28, color=RGBColor(0xCC, 0x00, 0x00), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "That's \u00a3749,400 per year in staff wages alone.", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")

    add_para(doc, "This is what it would cost to hire specialists to do what these 34 Claude Projects do for you. Not entry-level hires. Not interns. Actual experienced professionals who know what they're doing.")
    doc.add_paragraph("")
    add_para(doc, "And that's before you factor in employer National Insurance contributions (13.8%), pension auto-enrolment (3-5%), recruitment fees, software licenses, desk space, management time, sick days, and holiday cover.")
    doc.add_paragraph("")
    add_para(doc, "With those hidden costs added, the real figure is closer to \u00a3870,000-\u00a3910,000/year.", bold=True)
    doc.add_paragraph("")

    # ── YOUR ACTUAL COST ──
    add_para(doc, "YOUR ACTUAL COST TO RUN THIS ENTIRE SYSTEM:", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    add_table(doc,
        ['Tool', 'What It Does', 'Monthly Cost'],
        [
            ['Claude Pro', 'Powers all 34 AI projects', '\u00a315/month'],
            ['Make.com (Pro Plan)', 'Connects your tools and runs automations', '\u00a316/month'],
            ['Total', '', '\u00a331/month'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "\u00a331/month replaces \u00a362,450/month in staff wages.", bold=True, size=13, color=RGBColor(0x00, 0x66, 0x33), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "That's a 99.95% cost reduction.", bold=True, size=14, color=RGBColor(0x00, 0x66, 0x33), align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── FULL BREAKDOWN ──
    doc.add_page_break()
    add_heading(doc, "The Full Breakdown: All 34 Projects")
    add_para(doc, "Every project in this pack replaces a specialist role. Here's what you'd actually pay to hire each one.")
    doc.add_paragraph("")

    # Category 1
    add_heading(doc, "Category 1: Content Marketing", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['1. Content Strategy & Calendar', 'Content Strategist / Marketing Manager', '\u00a32,800'],
            ['2. Blog & Long-Form Content', 'Freelance Copywriter (long-form specialist)', '\u00a32,500'],
            ['3. Email Marketing Sequences', 'Email Marketing Specialist', '\u00a32,000'],
            ['4. Lead Magnets & Opt-In Content', 'Funnel Designer / Lead Gen Specialist', '\u00a32,500'],
            ['33. Community Management', 'Community Manager (membership/group)', '\u00a31,600'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a311,400/month (\u00a3136,800/year)", bold=True)
    add_para(doc, "A content marketing team of 5 specialists. Most business owners hire at least 2 of these roles before they hit \u00a3200K. You're getting all 5.")
    doc.add_paragraph("")

    # Category 2
    add_heading(doc, "Category 2: Social Media", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['5. Instagram Content Creator', 'Social Media Manager (Instagram specialist)', '\u00a31,600'],
            ['6. LinkedIn Thought Leadership', 'LinkedIn Content Strategist', '\u00a31,800'],
            ['7. YouTube Content & Scripts', 'YouTube Producer / Scriptwriter', '\u00a32,500'],
            ['8. Content Repurposing Engine', 'Content Repurposing Specialist', '\u00a31,800'],
            ['32. Podcast Strategy & Guest Appearances', 'Podcast Producer / PR Specialist', '\u00a32,000'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a39,700/month (\u00a3116,400/year)", bold=True)
    add_para(doc, "Most business owners outsource just ONE of these and spend \u00a31,500-\u00a32,500/month on it. This gives you all 5 platforms covered.")
    doc.add_paragraph("")

    # Category 3
    add_heading(doc, "Category 3: Sales & Lead Generation", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['9. Sales Landing Page Copy', 'Sales Copywriter (landing pages + sales pages)', '\u00a32,500'],
            ['10. Proposal & Pitch Writer', 'Business Development / Proposal Writer', '\u00a32,800'],
            ['11. DM & Outreach Sequences', 'Outreach Specialist / Lead Gen VA', '\u00a31,600'],
            ['12. Webinar & Workshop Planner', 'Event Planner / Webinar Producer', '\u00a31,400'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a38,300/month (\u00a399,600/year)", bold=True)
    add_para(doc, "This is the revenue engine. A sales copywriter alone runs \u00a32,500/month minimum for someone who actually converts. You're getting the entire sales support team.")
    doc.add_paragraph("")

    # Category 4
    doc.add_page_break()
    add_heading(doc, "Category 4: Client Fulfilment", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['13. Client Onboarding Assistant', 'Client Success Coordinator', '\u00a31,800'],
            ['14. Coaching Session Prep Notes', 'Executive Assistant / Session Admin', '\u00a31,400'],
            ['15. Client Reporting & Results', 'Reporting Analyst / Client Success Manager', '\u00a32,000'],
            ['16. Course & Programme Content', 'Instructional Designer / Course Creator', '\u00a31,800'],
            ['31. Client Offboarding & Retention', 'Client Retention Specialist', '\u00a31,800'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a38,800/month (\u00a3105,600/year)", bold=True)
    add_para(doc, "Client fulfilment roles are the ones you can never quite justify hiring but desperately need. Most businesses focus on onboarding and forget offboarding entirely. Retention is where hidden revenue lives.")
    doc.add_paragraph("")

    # Category 5
    add_heading(doc, "Category 5: Operations", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['17. SOP & Process Documentation', 'Operations Manager / Process Specialist', '\u00a33,000'],
            ['18. KPI & Business Reporting', 'Business Analyst / Data Analyst', '\u00a32,400'],
            ['19. Team Communication & Management', 'HR Coordinator / Team Manager', '\u00a31,400'],
            ['20. Finance & Invoicing Assistant', 'Bookkeeper / Finance Administrator', '\u00a31,300'],
            ['29. Finance & Strategic Reporting', 'Financial Controller / Part-Time CFO', '\u00a33,200'],
            ['34. Tech Stack & Automation Setup', 'Automation / CRM Specialist', '\u00a32,200'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a313,500/month (\u00a3162,000/year)", bold=True)
    add_para(doc, "Operations is where most business owners haemorrhage time. An operations manager alone costs \u00a336,000+/year. A part-time CFO runs \u00a33,200+. You're getting the entire ops department plus strategic finance and automation.")
    doc.add_paragraph("")

    # Category 6
    add_heading(doc, "Category 6: Strategy & Growth", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['21. Business Strategy & Planning', 'Business Consultant / Strategist', '\u00a32,500'],
            ['22. Brand Positioning Strategist', 'Brand Strategist / Positioning Expert', '\u00a32,400'],
            ['23. Partnership & Collaboration', 'Business Development Manager', '\u00a32,500'],
            ['24. Customer Research & Insights', 'Market Research Analyst', '\u00a32,400'],
            ['30. Offer Design & Pricing Strategy', 'Pricing Strategist / Revenue Consultant', '\u00a32,800'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a312,600/month (\u00a3151,200/year)", bold=True)
    add_para(doc, "Strategic roles are the most expensive to hire because you need senior experience. A pricing strategist alone charges \u00a3250-\u00a3500/day. You'd burn through \u00a32,800 in a week.")
    doc.add_paragraph("")

    # Category 7
    add_heading(doc, "Category 7: Admin & Organisation", level=2)
    add_table(doc,
        ['Project', 'Role It Replaces', 'Monthly Cost'],
        [
            ['25. Email & Communication Manager', 'Communications Manager / Admin', '\u00a32,200'],
            ['26. Meeting & Event Coordinator', 'Event Coordinator / Scheduling Admin', '\u00a31,200'],
            ['27. Legal & Compliance Drafts', 'Legal Administrator / Compliance Officer', '\u00a31,700'],
            ['28. Personal Productivity & Planning', 'Virtual Assistant / PA', '\u00a31,100'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Category Total: \u00a36,200/month (\u00a374,400/year)", bold=True)
    add_para(doc, "Admin roles feel small individually but add up fast. A decent VA costs \u00a31,100/month minimum, and that's just for general admin. Throw in legal, comms, and events and you're looking at \u00a36K+ before you've generated a single piece of content.")

    # ── SIDE BY SIDE ──
    doc.add_page_break()
    add_heading(doc, "The Numbers Side by Side")
    doc.add_paragraph("")
    add_table(doc,
        ['', 'Hiring Staff', 'Claude + Make System'],
        [
            ['Monthly Cost', '\u00a362,450', '\u00a331 (Claude Pro + Make Pro)'],
            ['Annual Cost', '\u00a3749,400', '\u00a3372'],
            ['Hidden Costs (NI, pension, recruitment)', '\u00a3120,000-\u00a3160,000', '\u00a30'],
            ['True Annual Cost', '\u00a3870,000-\u00a3910,000', '\u00a3372'],
            ['Availability', 'Business hours, minus holidays and sick days', '24/7, 365 days'],
            ['Ramp-Up Time', '3-6 months per hire', 'Same day'],
            ['Management Overhead', '5-10 hours/week managing team', 'Zero'],
            ['Recruitment Cost', '\u00a33,000-\u00a38,000 per role', 'Zero'],
            ['Risk of Bad Hire', 'Real (and expensive)', 'None'],
        ]
    )

    # ── WHAT THIS MEANS ──
    doc.add_paragraph("")
    add_heading(doc, "What This Actually Means For You")
    add_para(doc, "You're not going to hire 34 people. Nobody is. That's the whole point.")
    doc.add_paragraph("")

    add_para(doc, "What actually happens without this system:", bold=True)
    add_bullet(doc, "You do everything yourself (and work until 11 PM)")
    add_bullet(doc, "You hire 1-2 people and still can't cover everything")
    add_bullet(doc, "You outsource bits to freelancers and spend \u00a32,000-\u00a35,000/month for patchy coverage")
    add_bullet(doc, "The gaps in your business stay gaps")
    doc.add_paragraph("")

    add_para(doc, "What happens WITH this system:", bold=True)
    add_bullet(doc, "Every role is covered from day one")
    add_bullet(doc, "You work with Claude the way you'd brief a team member")
    add_bullet(doc, "Output quality matches or beats what a mid-level hire would produce")
    add_bullet(doc, "You keep the \u00a362,450/month in your pocket")

    # ── REAL COMPARISON ──
    doc.add_paragraph("")
    add_heading(doc, "The Real Comparison")
    add_para(doc, "Most business owners at \u00a350K-\u00a3150K revenue hire 1-3 of these roles and spend:")
    doc.add_paragraph("")
    add_table(doc,
        ['Typical Hire', 'Monthly Cost'],
        [
            ['Social Media Manager', '\u00a31,600'],
            ['VA / Admin Support', '\u00a31,100'],
            ['Copywriter (freelance)', '\u00a32,500'],
            ['Typical monthly spend', '\u00a35,200'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "That \u00a35,200/month gets you 3 out of 34 roles covered. Partially.")
    doc.add_paragraph("")
    add_para(doc, "This system gives you all 34 for \u00a331/month.", bold=True, size=13)
    doc.add_paragraph("")
    add_para(doc, "That's a 99.95% cost reduction.", bold=True, size=14, color=RGBColor(0x00, 0x66, 0x33))

    # ── CAVEATS ──
    doc.add_paragraph("")
    add_heading(doc, "Important Caveats")
    add_para(doc, "This system doesn't replace:", bold=True)
    add_bullet(doc, "Final decision-making (strategy still needs your brain)")
    add_bullet(doc, "Client delivery (you still do the coaching and consulting)")
    add_bullet(doc, "Relationship building (human connection can't be automated)")
    add_bullet(doc, "Quality review (always review AI output before publishing)")
    doc.add_paragraph("")

    add_para(doc, "What it DOES replace:", bold=True)
    add_bullet(doc, "The first draft of everything")
    add_bullet(doc, "Research, planning, and preparation")
    add_bullet(doc, "Repetitive admin and operational tasks")
    add_bullet(doc, "Content creation across all channels")
    add_bullet(doc, "Documentation, reporting, and analysis")
    add_bullet(doc, "The 60-70% of your week currently spent on operational work")

    # ── BOTTOM LINE ──
    doc.add_paragraph("")
    add_heading(doc, "Bottom Line")
    add_para(doc, "34 specialist roles. \u00a362,450/month in staff wages. Replaced by a \u00a331/month system you set up once and use forever.", bold=True, size=13)
    doc.add_paragraph("")
    add_para(doc, "The question isn't whether you can afford to use Claude Projects.", size=12)
    add_para(doc, "The question is whether you can afford not to.", bold=True, size=13)
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_para(doc, "All salary data based on 2025-2026 UK specialist/freelancer market rates for experienced professionals (3-5+ years). Costs reflect fractional or part-time engagement typical for small businesses and solo operators. Sources include Robert Half UK Salary Guide, Glassdoor UK, Reed.co.uk, and YunoJuno Freelancer Rates Report.", size=8, color=RGBColor(0x88, 0x88, 0x88))

    # Save
    onboarding_dir = os.path.join(BASE_DIR, "00-Onboarding")
    os.makedirs(onboarding_dir, exist_ok=True)
    filepath = os.path.join(onboarding_dir, "00-COST-SAVINGS-BREAKDOWN.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == "__main__":
    main()
