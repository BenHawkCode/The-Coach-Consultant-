"""
Generate Claude Code Guides for Claude Projects Client Pack.
Creates the 08-Claude-Code-Guides folder with ready-to-use prompts
for managing, updating, and maintaining Claude Projects at scale
using Claude Code (the CLI/desktop tool).
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "Claude-Projects-Client-Pack")
GUIDES_DIR = os.path.join(BASE_DIR, "08-Claude-Code-Guides")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. Unauthorised sharing, copying, reproduction, or sale is prohibited without written permission. Legal action will be taken for infringements. Personal use is allowed for active members or opt-ins only."
FONT_NAME = "Poppins"


# ─── Styling helpers (matched to existing pack styling) ───

def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for style_name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = FONT_NAME


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = FONT_NAME
    return p


def add_checkbox(doc, text):
    doc.add_paragraph(f"\u2610  {text}", style='List Bullet')


def add_code_block(doc, text):
    """Add a monospaced code block for copy-paste prompts."""
    doc.add_paragraph("")
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    return table


def save_doc(doc, filename):
    os.makedirs(GUIDES_DIR, exist_ok=True)
    filepath = os.path.join(GUIDES_DIR, filename)
    doc.save(filepath)
    print(f"  Created: {filepath}")


def new_doc():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)
    return doc


# ─── Document 0: What Is Claude Code & How to Use It ───

def create_intro_guide():
    doc = new_doc()

    add_heading(doc, "What Is Claude Code and Why You Need It")
    add_para(doc, "How to use Claude's desktop tool to manage, update, and maintain all of your Claude Projects at scale.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "What Is Claude Code?", level=2)
    add_para(doc, "Claude Code is a separate tool from the Claude website (claude.ai). While the website is where you chat with Claude inside your Projects, Claude Code is a desktop application that lets you give Claude bigger, multi-step tasks and have it work through them systematically.")
    doc.add_paragraph("")
    add_para(doc, "Think of it this way:", bold=True)
    add_bullet(doc, "Claude.ai (the website) = Your team members. You chat with them one project at a time")
    add_bullet(doc, "Claude Code (the desktop tool) = Your operations manager. It can work across multiple projects, run through checklists, and handle bulk updates")
    doc.add_paragraph("")

    add_heading(doc, "When to Use Claude Code vs Claude.ai", level=2)
    doc.add_paragraph("")
    add_table(doc,
        ['Task', 'Use Claude.ai', 'Use Claude Code'],
        [
            ['Write an Instagram caption', 'Yes', 'No'],
            ['Create a sales email', 'Yes', 'No'],
            ['Update brand files across all 34 projects', 'No (too slow)', 'Yes'],
            ['Audit all project instructions for consistency', 'No (too slow)', 'Yes'],
            ['Change your CTA across every project', 'No (too slow)', 'Yes'],
            ['Generate a content calendar', 'Yes', 'No'],
            ['Bulk update pricing in all projects', 'No (too slow)', 'Yes'],
            ['Back up all your project configurations', 'No', 'Yes'],
            ['Run a quarterly project health check', 'No (too slow)', 'Yes'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "The simple rule: if you need to do the same thing across multiple projects, use Claude Code. If you are working inside one project on one task, use Claude.ai.", bold=True)
    doc.add_paragraph("")

    add_heading(doc, "How to Get Claude Code", level=2)
    add_para(doc, "Step 1: Go to claude.ai/download (or search 'Claude Code download')")
    add_para(doc, "Step 2: Download the desktop application for your operating system (Mac or Windows)")
    add_para(doc, "Step 3: Install it and sign in with the same account you use for Claude.ai")
    add_para(doc, "Step 4: You are ready to go")
    doc.add_paragraph("")
    add_para(doc, "Claude Code uses the same subscription as your Claude.ai account. No extra cost.")
    doc.add_paragraph("")

    add_heading(doc, "How to Use the Prompts in This Folder", level=2)
    add_para(doc, "Each document in this folder contains a ready-to-use prompt. Here is how to use them:")
    doc.add_paragraph("")
    add_para(doc, "Step 1: Open the document and read what it does")
    add_para(doc, "Step 2: Copy the prompt (the text in the code block)")
    add_para(doc, "Step 3: Replace all the [PLACEHOLDER] fields with your actual details")
    add_para(doc, "Step 4: Open Claude Code on your desktop")
    add_para(doc, "Step 5: Paste the prompt and press Enter")
    add_para(doc, "Step 6: Let Claude work through it step by step. It will tell you what it is doing as it goes")
    doc.add_paragraph("")
    add_para(doc, "These prompts are designed to be thorough. Claude Code will work through each step systematically and give you a summary when it is done. You do not need to babysit it.", bold=True)
    doc.add_paragraph("")

    add_heading(doc, "Important Notes", level=2)
    add_bullet(doc, "Always review what Claude Code has done before moving on. Check the summary it gives you")
    add_bullet(doc, "Keep a backup of your current project instructions before running bulk updates")
    add_bullet(doc, "Claude Code cannot upload files to claude.ai for you. It will tell you which files to upload and where. You still do the clicking")
    add_bullet(doc, "These prompts work with Claude Code's ability to read your local files and generate instructions. The actual project updates on claude.ai are done by you following Claude's guidance")

    save_doc(doc, "00-WHAT-IS-CLAUDE-CODE.docx")


# ─── Document 1: Full Rebrand Update ───

def create_rebrand_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Full Rebrand Update Across All Projects")
    add_para(doc, "Use this when you have rebranded and need to update every single Claude Project with new brand files, names, colours, and messaging.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You have changed your business name")
    add_bullet(doc, "You have updated your brand identity (colours, fonts, logo)")
    add_bullet(doc, "You have rewritten your brand voice or tone guidelines")
    add_bullet(doc, "You have updated your messaging framework or positioning")
    add_bullet(doc, "Any combination of the above")
    doc.add_paragraph("")

    add_heading(doc, "Before You Start", level=2)
    add_checkbox(doc, "Have your new brand files ready in one folder on your computer")
    add_checkbox(doc, "Know the names of the old files you need to remove")
    add_checkbox(doc, "Have a list of all your active Claude Projects")
    add_checkbox(doc, "Back up your current project instructions (copy them into a document)")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)
    add_para(doc, "Copy everything below, replace the [PLACEHOLDER] fields, and paste it into Claude Code:")

    add_code_block(doc, """I have completed a rebrand and need to update all of my Claude Projects with new brand files and settings.
Please work through every project listed below systematically.
Do not move to the next project until the current one is fully updated.

---

MY BRAND: [YOUR BRAND NAME]

NEW BRAND FILES TO ADD:
(These are all located at: [folder path / Drive folder / Desktop folder])

- [File 1 - e.g. Brand Guidelines.pdf]
- [File 2 - e.g. Brand Voice & Tone Guide.docx]
- [File 3 - e.g. Logo Assets.zip]
- [File 4 - e.g. Messaging Framework.docx]
- [Add or remove as needed]

FILES TO REMOVE FROM EACH PROJECT:
(Old versions to delete before uploading new ones)

- [Old File 1 name]
- [Old File 2 name]
- [Add or remove as needed]

MY PROJECTS TO UPDATE:
- [Project 1]
- [Project 2]
- [Project 3]
- [Project 4]
- [Add as needed]

---

FOR EACH PROJECT, FOLLOW THESE STEPS IN ORDER:

1. Open claude.ai and navigate to my Projects list
2. Open the project
3. Go to the project files or settings section
4. Remove every old brand file listed above if present
5. Upload every new brand file listed above
6. Check the project system prompt for any outdated references to old brand names, colours, tone descriptions, or positioning language
7. If outdated references are found, update them to reflect my new brand
8. Confirm all files are present and the system prompt is updated
9. Log the project as complete before moving on

---

COMPLETION CHECK:

Once all projects are updated, provide me with a summary listing:
- Each project name
- Files removed
- Files added
- Whether the system prompt was updated and what changed

Do not stop until all projects are complete and the summary is delivered.""")

    add_heading(doc, "What Claude Code Will Do", level=2)
    add_para(doc, "Claude Code will work through your project list one by one. For each project it will:")
    add_bullet(doc, "Tell you which old files to remove")
    add_bullet(doc, "Tell you which new files to upload")
    add_bullet(doc, "Review your system prompt and flag any outdated brand references")
    add_bullet(doc, "Generate updated instruction text where needed")
    add_bullet(doc, "Give you a completion summary at the end")
    doc.add_paragraph("")
    add_para(doc, "You follow along in claude.ai making the changes as Claude Code guides you through each project.")

    save_doc(doc, "01-FULL-REBRAND-UPDATE.docx")


# ─── Document 2: Bulk CTA Update ───

def create_cta_update_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Update Your CTA Across All Projects")
    add_para(doc, "Use this when your call-to-action has changed. New website URL, new booking link, new lead magnet, or new offer.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You have changed your website URL")
    add_bullet(doc, "You have a new booking link (e.g. moved from Calendly to TidyCal)")
    add_bullet(doc, "You have launched a new lead magnet and want all projects to reference it")
    add_bullet(doc, "You have changed your main offer and the CTA needs updating")
    add_bullet(doc, "You have updated your email address or contact method")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I need to update my call-to-action across all of my Claude Projects.
Work through every project systematically and do not skip any.

---

OLD CTA DETAILS (what needs replacing):
- Old website: [e.g. www.oldsite.com]
- Old booking link: [e.g. calendly.com/oldname]
- Old lead magnet: [e.g. "Download my free checklist"]
- Old offer name: [e.g. "The Starter Package"]
- Old email: [e.g. old@email.com]
(Remove any lines that do not apply)

NEW CTA DETAILS (what to replace them with):
- New website: [e.g. www.newsite.com]
- New booking link: [e.g. tidycal.com/newname]
- New lead magnet: [e.g. "Grab my free Growth Playbook"]
- New offer name: [e.g. "The Growth Accelerator"]
- New email: [e.g. new@email.com]
(Remove any lines that do not apply)

MY PROJECTS TO UPDATE:
- [Project 1]
- [Project 2]
- [Project 3]
- [Add all your projects]

---

FOR EACH PROJECT:

1. Open the project's custom instructions
2. Search for every instance of the old CTA details listed above
3. Replace each instance with the corresponding new CTA detail
4. Check knowledge files for any documents that reference old CTA details and flag them
5. Log what was changed before moving to the next project

---

COMPLETION SUMMARY:

Provide a table showing:
- Project name
- Number of CTA references found
- What was changed
- Any knowledge files that need manual updating

Do not stop until every project is checked and updated.""")

    save_doc(doc, "02-BULK-CTA-UPDATE.docx")


# ─── Document 3: Pricing & Offer Update ───

def create_pricing_update_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Update Pricing and Offers Across All Projects")
    add_para(doc, "Use this when your pricing changes, you launch a new offer, retire an old one, or restructure your offer stack.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You have changed the price of an existing offer")
    add_bullet(doc, "You have launched a new offer or programme")
    add_bullet(doc, "You have retired or discontinued an offer")
    add_bullet(doc, "You have restructured your offer stack (e.g. added tiers)")
    add_bullet(doc, "You have changed payment terms or plan options")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I have updated my offers and pricing and need every Claude Project to reflect the changes.
Work through each project systematically.

---

MY UPDATED OFFER STACK:

Offer 1: [Offer Name]
- Price: [e.g. £997 or £297/month x 4]
- What it includes: [brief description]
- Who it is for: [target audience for this offer]
- CTA: [e.g. "Book a discovery call at www.yoursite.com/call"]

Offer 2: [Offer Name]
- Price: [price]
- What it includes: [brief description]
- Who it is for: [target audience]
- CTA: [CTA for this offer]

Offer 3: [Offer Name]
- Price: [price]
- What it includes: [brief description]
- Who it is for: [target audience]
- CTA: [CTA for this offer]

(Add or remove offers as needed)

OFFERS THAT HAVE BEEN RETIRED (remove all references):
- [Old offer name 1]
- [Old offer name 2]
(Remove this section if not applicable)

MY PROJECTS TO UPDATE:
- [Project 1]
- [Project 2]
- [Add all your projects]

---

FOR EACH PROJECT:

1. Review the custom instructions for any pricing, offer names, or offer descriptions
2. Update all references to match my new offer stack above
3. Remove any references to retired offers
4. Check if the project's knowledge files mention old pricing (flag for manual update)
5. Ensure the CTA matches the correct offer for that project's purpose
6. Log all changes before moving to the next project

---

COMPLETION SUMMARY:

For each project list:
- What pricing/offer references were found
- What was updated
- What was removed
- Any knowledge files flagged for manual update""")

    save_doc(doc, "03-PRICING-AND-OFFER-UPDATE.docx")


# ─── Document 4: Target Audience Update ───

def create_audience_update_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Update Target Audience Across All Projects")
    add_para(doc, "Use this when your ideal client has shifted, you are niching down further, or your audience demographics and pain points have changed.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You have narrowed your niche (e.g. from 'coaches' to 'fitness coaches')")
    add_bullet(doc, "Your ideal client profile has changed (different revenue level, different stage)")
    add_bullet(doc, "You have updated your audience's key pain points based on new research")
    add_bullet(doc, "You are pivoting to serve a different market segment")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I have refined my target audience and need all of my Claude Projects updated to reflect this change.
Every project should speak to and create content for my updated ideal client.

---

MY UPDATED TARGET AUDIENCE:

Who I help: [e.g. "Female fitness coaches in the UK"]
Their stage: [e.g. "Making £50K-£150K, 1-3 years in business"]
Their biggest pain points:
- [Pain point 1 - e.g. "Stuck trading time for money with 1-1 clients"]
- [Pain point 2 - e.g. "No systems, doing everything manually"]
- [Pain point 3 - e.g. "Know they need to create content but never have time"]

What they want:
- [Desire 1 - e.g. "A scalable offer that does not require their time for every sale"]
- [Desire 2 - e.g. "Systems that run without them"]
- [Desire 3 - e.g. "To be seen as the go-to expert in their space"]

Language they use: [e.g. "They say 'I am always busy but not making more money', 'I know I should be doing more online', 'I do not know where to start with scaling'"]

OLD AUDIENCE REFERENCES TO REPLACE:
- Old niche: [e.g. "coaches and consultants"]
- Old descriptors: [e.g. "business owners", "entrepreneurs"]
- Old pain points to remove: [e.g. "struggling to get started"]

MY PROJECTS TO UPDATE:
- [Project 1]
- [Project 2]
- [Add all your projects]

---

FOR EACH PROJECT:

1. Review the custom instructions for audience references
2. Update all niche descriptors, audience descriptions, and pain points
3. Update any example prompts or suggested content topics that reference the old audience
4. Check knowledge files for audience-related documents and flag for manual update
5. Ensure the tone and language matches what my new audience responds to
6. Log all changes

---

COMPLETION SUMMARY:

For each project:
- Audience references found
- What was updated
- Knowledge files flagged for manual update""")

    save_doc(doc, "04-TARGET-AUDIENCE-UPDATE.docx")


# ─── Document 5: Quarterly Project Health Check ───

def create_quarterly_audit_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Quarterly Project Health Check")
    add_para(doc, "Run this every 3 months to audit all your Claude Projects for consistency, outdated information, and optimisation opportunities.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "Every quarter (put it in your calendar)")
    add_bullet(doc, "After any major business change")
    add_bullet(doc, "When you notice projects producing inconsistent or outdated output")
    add_bullet(doc, "When you add new projects and want to check alignment with existing ones")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """Run a full health check on all of my Claude Projects.
Audit each one for accuracy, consistency, and optimisation opportunities.

---

MY CURRENT BUSINESS DETAILS:

Brand name: [YOUR BRAND NAME]
Website: [YOUR WEBSITE]
Main CTA: [YOUR CURRENT CTA]
Current offers: [LIST YOUR CURRENT OFFERS AND PRICES]
Target audience: [YOUR CURRENT TARGET AUDIENCE DESCRIPTION]
Brand voice: [e.g. "Direct, conversational, no corporate jargon, British English"]

MY PROJECTS TO AUDIT:
- [Project 1]
- [Project 2]
- [Add all your projects]

---

FOR EACH PROJECT, CHECK:

1. ACCURACY
   - Does the brand name match exactly?
   - Is the website URL correct?
   - Are the offers and pricing current?
   - Is the target audience description up to date?
   - Are all CTAs pointing to the right place?

2. CONSISTENCY
   - Does the brand voice description match across all projects?
   - Are the same forbidden phrases listed in every writing project?
   - Is the tone description consistent?
   - Do all projects reference the same offers and audience?

3. KNOWLEDGE FILES
   - Are there outdated files that should be replaced?
   - Are there files missing that should be uploaded?
   - Are file names clear and descriptive?

4. OPTIMISATION
   - Are instructions clear and specific enough?
   - Could any instructions be improved based on how the project is actually used?
   - Are there any redundant or contradictory instructions?
   - Are example prompts relevant and useful?

---

DELIVER THREE THINGS:

1. A project-by-project status report (Healthy / Needs Update / Needs Attention)
2. A prioritised list of all issues found, grouped by severity
3. Updated instruction text for any project that needs changes (ready to copy-paste)""")

    save_doc(doc, "05-QUARTERLY-HEALTH-CHECK.docx")


# ─── Document 6: New Knowledge File Rollout ───

def create_knowledge_rollout_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Add a New Knowledge File to Multiple Projects")
    add_para(doc, "Use this when you have created a new document (brand voice guide, case study, SOP, etc.) that needs to be added to several projects at once.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You have created a new brand voice guide and want it in every writing project")
    add_bullet(doc, "You have a new case study that multiple projects should reference")
    add_bullet(doc, "You have written an updated SOP that applies to several projects")
    add_bullet(doc, "You have new client research that should inform content and strategy projects")
    add_bullet(doc, "You have updated your competitor analysis")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I have a new knowledge file that needs to be added to multiple Claude Projects.
Help me identify which projects need it and guide me through uploading it.

---

NEW FILE DETAILS:

File name: [e.g. "Brand-Voice-Guide-v2.pdf"]
File location: [e.g. "Desktop/Brand Files/"]
What it contains: [e.g. "Updated brand voice guidelines including tone, forbidden phrases, example content, and voice samples"]
File it replaces (if any): [e.g. "Brand-Voice-Guide-v1.pdf" or "None - this is new"]

MY PROJECTS:
- [Project 1]
- [Project 2]
- [Add all your projects]

---

FOR EACH PROJECT:

1. Determine whether this file is relevant to the project's purpose
2. If relevant, check whether the project already has the old version
3. Tell me to remove the old version (if applicable)
4. Tell me to upload the new version
5. Check if the project's custom instructions need updating to reference the new file
6. If instructions need updating, provide the updated text

---

COMPLETION SUMMARY:

Provide a table showing:
- Project name
- Whether the file was added (Yes / No / Not Relevant)
- Old file removed (Yes / No / N/A)
- Instructions updated (Yes / No)
- Notes""")

    save_doc(doc, "06-NEW-KNOWLEDGE-FILE-ROLLOUT.docx")


# ─── Document 7: Voice & Tone Recalibration ───

def create_voice_recalibration_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Voice and Tone Recalibration")
    add_para(doc, "Use this when your voice or tone has evolved and your projects are producing content that does not sound like you any more.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "Your writing style has evolved and Claude's output feels outdated")
    add_bullet(doc, "You have new content samples that better represent how you communicate now")
    add_bullet(doc, "You are getting consistent feedback that AI content does not sound like you")
    add_bullet(doc, "You want to tighten up voice consistency across all projects")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """My brand voice has evolved and I need to recalibrate all of my writing-focused Claude Projects.
Analyse my new voice samples and update every relevant project.

---

MY NEW VOICE SAMPLES:
(Located at: [folder path])

- [File 1 - e.g. "Recent-Instagram-Captions-2025.docx"]
- [File 2 - e.g. "Latest-Email-Sequence.docx"]
- [File 3 - e.g. "Voice-Memo-Transcription.txt"]
- [File 4 - e.g. "Latest-YouTube-Script.docx"]

WHAT HAS CHANGED ABOUT MY VOICE:
- [e.g. "I am more direct now, less wordy"]
- [e.g. "I use more stories and fewer bullet points"]
- [e.g. "My tone is warmer and less 'teacher mode'"]
- [e.g. "I have new phrases I use regularly: [list them]"]
- [e.g. "I have stopped using these phrases: [list them]"]

OLD VOICE FILES TO REMOVE:
- [Old File 1]
- [Old File 2]

MY WRITING PROJECTS TO UPDATE:
- [Project 1 - e.g. Instagram Content Creator]
- [Project 2 - e.g. Email Marketing & Sequences]
- [Project 3 - e.g. Blog & Long-Form Content]
- [Add all writing-related projects]

---

FOR EACH WRITING PROJECT:

1. Remove old voice sample files
2. Upload new voice sample files
3. Review the voice rules in the custom instructions
4. Update forbidden phrases list if my "do not use" list has changed
5. Update recommended phrases list if I have new signature phrases
6. Update any tone descriptors in the instructions
7. Ensure the voice rules section is consistent across all writing projects
8. Log every change

---

COMPLETION SUMMARY:

For each project:
- Old voice files removed
- New voice files uploaded
- Instruction changes made (exact text of what changed)
- Consistency check result (matches other projects: Yes/No)""")

    save_doc(doc, "07-VOICE-AND-TONE-RECALIBRATION.docx")


# ─── Document 8: Seasonal Campaign Setup ───

def create_campaign_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Seasonal or Launch Campaign Setup")
    add_para(doc, "Use this to temporarily update multiple projects for a specific campaign, launch, or seasonal push, then revert them afterwards.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You are running a launch and need all content to reference the offer")
    add_bullet(doc, "Black Friday, New Year, or seasonal promotions")
    add_bullet(doc, "You are doing a challenge, webinar series, or live event")
    add_bullet(doc, "Any time-limited campaign where multiple projects need temporary updates")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I am running a campaign and need to temporarily update several Claude Projects to support it.
After the campaign ends I will need to revert these changes, so document everything clearly.

---

CAMPAIGN DETAILS:

Campaign name: [e.g. "January Reset Launch"]
Start date: [date]
End date: [date]
Campaign offer: [e.g. "The Growth Accelerator - £497 launch price (normally £997)"]
Campaign CTA: [e.g. "Join the January Reset at www.yoursite.com/january"]
Campaign hashtag (if any): [e.g. #JanuaryReset]
Key messaging: [e.g. "New year, new systems. Build your AI-powered business in 30 days"]

PROJECTS TO UPDATE FOR THIS CAMPAIGN:
- [Project 1 - e.g. Instagram Content Creator]
- [Project 2 - e.g. Email Marketing & Sequences]
- [Project 3 - e.g. DM & Outreach Sequences]
- [Project 4 - e.g. Sales Landing Page Copy]
- [Add as needed]

---

FOR EACH PROJECT:

1. Save a copy of the current custom instructions (label it "PRE-CAMPAIGN BACKUP")
2. Add a campaign context section to the top of the instructions:
   "ACTIVE CAMPAIGN: [campaign name]. All content should support this campaign where relevant. Campaign runs [start] to [end]. Primary CTA: [campaign CTA]. Key messaging: [key messaging]"
3. Update any existing CTAs to reference the campaign offer during the campaign period
4. Note exactly what was changed so it can be reverted

---

DELIVER TWO THINGS:

1. A project-by-project change log showing exactly what was modified
2. A "REVERT GUIDE" I can use after the campaign to put everything back to normal,
   with the exact original text for each instruction that was changed""")

    save_doc(doc, "08-SEASONAL-CAMPAIGN-SETUP.docx")


# ─── Document 9: Project Backup & Export ───

def create_backup_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Back Up All Project Configurations")
    add_para(doc, "Use this to create a complete backup of all your project instructions, knowledge file lists, and settings. Run this before any major changes.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "Before running any bulk update from this folder")
    add_bullet(doc, "Monthly or quarterly as a safety net")
    add_bullet(doc, "Before letting a team member or VA make changes to your projects")
    add_bullet(doc, "When you want a snapshot of your current setup to share or reference")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """Create a complete backup document of all my Claude Project configurations.
I need this as a reference and safety net before making changes.

---

MY PROJECTS TO BACK UP:
- [Project 1]
- [Project 2]
- [Project 3]
- [Add all your projects]

---

FOR EACH PROJECT, DOCUMENT:

1. Project name (exactly as it appears in Claude)
2. Full custom instructions (copy every word)
3. List of all knowledge files currently uploaded (file names and descriptions)
4. Any notes about how this project is configured

---

FORMAT:

Create a single document with a clear table of contents.
Each project gets its own section with:
- Project Name (as heading)
- Custom Instructions (in a code block for easy copy-paste)
- Knowledge Files (as a numbered list)
- Configuration Notes (any relevant observations)

Add a timestamp at the top: "Backup created: [today's date]"

This document should be comprehensive enough that I could rebuild any project
from scratch using only this backup.""")

    save_doc(doc, "09-PROJECT-BACKUP-AND-EXPORT.docx")


# ─── Document 10: New Project Setup from Template ───

def create_new_project_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Set Up a New Project from Your Pack Template")
    add_para(doc, "Use this when you are setting up a new project from your Client Pack and want Claude Code to guide you through every step including customisation.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You are setting up a project you have not used before")
    add_bullet(doc, "You want Claude Code to help you customise the instructions for your specific business")
    add_bullet(doc, "You want a guided walkthrough rather than doing it from the document alone")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I want to set up a new Claude Project from my Client Pack.
Guide me through the complete setup step by step.

---

PROJECT I AM SETTING UP:
[e.g. "Project 5: Instagram Content Creator"]

MY BUSINESS DETAILS:
- Brand name: [YOUR BRAND NAME]
- Niche: [YOUR NICHE]
- Target audience: [YOUR AUDIENCE]
- Brand voice: [e.g. "Direct, conversational, British English, no corporate jargon"]
- Website: [YOUR WEBSITE]
- Main CTA: [YOUR CTA]
- Main offer: [YOUR MAIN OFFER]

---

PLEASE:

1. Read the project template from my Client Pack
2. Replace all [PLACEHOLDER] fields with my actual business details
3. Give me the finalised custom instructions ready to paste into Claude
4. Tell me exactly which knowledge files to upload for this project
5. Suggest 3 starter prompts I can use to test the project
6. Flag any customisations I should consider based on my niche

Make it completely ready to use. I should be able to copy your output, paste it
into Claude, upload the files you list, and start working immediately.""")

    save_doc(doc, "10-NEW-PROJECT-SETUP-GUIDED.docx")


# ─── Document 11: Compliance & Disclaimer Update ───

def create_compliance_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Update Compliance and Disclaimers Across Projects")
    add_para(doc, "Use this when legal requirements change, you update your terms, or you need to add disclaimers to projects that reference financial, health, or legal topics.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "GDPR or data protection rules change")
    add_bullet(doc, "You update your terms of service or privacy policy")
    add_bullet(doc, "You need to add income or results disclaimers to marketing projects")
    add_bullet(doc, "Your industry regulator issues new guidelines")
    add_bullet(doc, "You want to ensure all projects have appropriate disclaimers")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """I need to update compliance and disclaimer language across my Claude Projects.
Audit every project and ensure they all meet my current requirements.

---

MY COMPLIANCE REQUIREMENTS:

1. Financial claims: [e.g. "All income references must include 'Results may vary. These figures are not guaranteed.'"]
2. Testimonials: [e.g. "All testimonials must note 'Individual results may vary'"]
3. Legal disclaimers: [e.g. "Project 27 (Legal & Compliance) must state 'Drafts only, not legal advice'"]
4. Data protection: [e.g. "No project should instruct Claude to store or reference personal client data beyond the conversation"]
5. Health claims: [e.g. "No project should make specific health outcome promises"]
(Add or remove as needed for your industry)

MY PROJECTS TO AUDIT:
- [Project 1]
- [Project 2]
- [Add all your projects]

---

FOR EACH PROJECT:

1. Review custom instructions for any claims, promises, or language that needs a disclaimer
2. Check if appropriate disclaimers are already present
3. Add missing disclaimers to the instructions where needed
4. Flag any knowledge files that contain claims needing disclaimers
5. Ensure the project does not instruct Claude to make claims beyond what is allowed

---

DELIVER:

1. An audit report showing which projects needed changes
2. Updated instruction text for any project that needed disclaimers added
3. A list of knowledge files flagged for manual review""")

    save_doc(doc, "11-COMPLIANCE-AND-DISCLAIMER-UPDATE.docx")


# ─── Document 12: Team Handover Pack ───

def create_team_handover_prompt():
    doc = new_doc()

    add_heading(doc, "Prompt: Create a Team or VA Handover for Your Projects")
    add_para(doc, "Use this to generate a complete handover document so a team member or VA can use your Claude Projects without needing you to explain everything.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "When to Use This", level=2)
    add_bullet(doc, "You are hiring a VA and want them to use your Claude Projects")
    add_bullet(doc, "A team member is taking over responsibility for certain projects")
    add_bullet(doc, "You want to document how to use each project so you are not the bottleneck")
    add_bullet(doc, "You are going on holiday and someone else needs to keep things running")
    doc.add_paragraph("")

    add_heading(doc, "The Prompt (Copy and Customise)", level=2)

    add_code_block(doc, """Create a complete handover guide for my Claude Projects.
This will be used by [a VA / team member / contractor] who needs to
use these projects without my help.

---

WHO IS THIS FOR:
- Name/role: [e.g. "Sarah, my VA"]
- Their Claude experience: [e.g. "Has used ChatGPT but never Claude Projects"]
- What they need access to: [e.g. "Content creation projects only" or "All projects"]

PROJECTS THEY WILL USE:
- [Project 1]
- [Project 2]
- [Add relevant projects]

---

FOR EACH PROJECT, CREATE:

1. A plain-English summary of what this project does (2-3 sentences)
2. Step-by-step instructions for opening the project and starting a chat
3. The 3-5 most common tasks they will use this project for
4. Example prompts for each common task (ready to copy-paste)
5. Common mistakes to avoid with this project
6. How to tell if the output quality is good or needs refining
7. What to do if they get stuck (who to ask, how to troubleshoot)

---

FORMAT:

Create a single handover document with:
- A table of contents
- One section per project
- An FAQ section at the end covering common Claude questions
- A quick-start checklist they can follow on day one

Write it assuming they are intelligent but have never used Claude Projects before.
Keep it practical and jargon-free.""")

    save_doc(doc, "12-TEAM-VA-HANDOVER-PACK.docx")


# ─── Generate everything ───

def main():
    print("Generating Claude Code Guides...\n")

    os.makedirs(GUIDES_DIR, exist_ok=True)

    create_intro_guide()
    create_rebrand_prompt()
    create_cta_update_prompt()
    create_pricing_update_prompt()
    create_audience_update_prompt()
    create_quarterly_audit_prompt()
    create_knowledge_rollout_prompt()
    create_voice_recalibration_prompt()
    create_campaign_prompt()
    create_backup_prompt()
    create_new_project_prompt()
    create_compliance_prompt()
    create_team_handover_prompt()

    print(f"\nDone! 13 Claude Code guide documents created in: {GUIDES_DIR}")


if __name__ == "__main__":
    main()
