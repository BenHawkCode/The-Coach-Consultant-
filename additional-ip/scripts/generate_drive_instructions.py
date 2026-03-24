"""
Generate Google Drive Setup Instructions for Claude Projects Client Pack.
Creates a document explaining how to add the pack to their Drive
and access it as a live-updated resource.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "Claude-Projects-Client-Pack")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. Unauthorised sharing, copying, reproduction, or sale is prohibited without written permission. Legal action will be taken for infringements. Personal use is allowed for active members or opt-ins only."
FONT_NAME = "Poppins"


def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for style_name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = FONT_NAME


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = FONT_NAME
    return p


def main():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)

    # ── Title ──
    add_heading(doc, "How to Access Your Client Pack")
    add_para(doc, "Important: Read this first before opening any other documents.", size=13, color=RGBColor(0xCC, 0x44, 0x44))
    doc.add_paragraph("")

    # ── What this is ──
    add_heading(doc, "What Is This Folder?", level=2)
    add_para(doc, "This Google Drive folder contains your complete Claude Projects Client Pack. It is a live resource that gets updated over time with new templates, improvements, and additional guides.")
    doc.add_paragraph("")
    add_para(doc, "Because this is a live resource, it is important that you access it correctly so you always have the latest version of everything.", bold=True)
    doc.add_paragraph("")

    # ── Add shortcut ──
    add_heading(doc, "Step 1: Add a Shortcut to Your Drive", level=2)
    add_para(doc, "Do NOT download these files. Instead, add a shortcut so the folder appears in your own Google Drive and stays synced with any updates we make.", bold=True)
    doc.add_paragraph("")

    add_para(doc, "On Desktop (Browser):", bold=True)
    add_bullet(doc, "Open this folder in Google Drive")
    add_bullet(doc, "Look at the folder name at the top of the page")
    add_bullet(doc, "Click the small dropdown arrow next to the folder name")
    add_bullet(doc, "Select 'Organise' (or 'Add shortcut to Drive' depending on your version)")
    add_bullet(doc, "Choose 'My Drive' or a specific folder where you want the shortcut")
    add_bullet(doc, "Click 'Add'")
    doc.add_paragraph("")

    add_para(doc, "Alternative method:", bold=True)
    add_bullet(doc, "Right-click on this folder in Google Drive")
    add_bullet(doc, "Select 'Organise' then 'Add shortcut'")
    add_bullet(doc, "Choose where you want it in your Drive")
    add_bullet(doc, "Click 'Add'")
    doc.add_paragraph("")

    add_para(doc, "On Mobile (Google Drive App):", bold=True)
    add_bullet(doc, "Open the folder in the Google Drive app")
    add_bullet(doc, "Tap the three dots menu (top right)")
    add_bullet(doc, "Tap 'Add shortcut to Drive'")
    add_bullet(doc, "Choose 'My Drive' and tap 'Add'")
    doc.add_paragraph("")

    add_para(doc, "Once you have added the shortcut, this folder will appear in your Google Drive as if it is your own folder. You can access it from your Drive home screen, and it will always show the latest files.", bold=True)
    doc.add_paragraph("")

    # ── Do not download ──
    add_heading(doc, "Step 2: Do NOT Download the Files", level=2)
    add_para(doc, "This is a live resource. We update it regularly with:", bold=True)
    add_bullet(doc, "Improved project templates based on client feedback")
    add_bullet(doc, "New projects as we build them")
    add_bullet(doc, "Updated onboarding documents")
    add_bullet(doc, "New Claude Code guides for managing your projects")
    add_bullet(doc, "Bug fixes and improvements to existing documents")
    doc.add_paragraph("")

    add_para(doc, "If you download the files, you get a frozen snapshot that will become outdated. By accessing them through the Drive shortcut, you always see the latest version automatically.")
    doc.add_paragraph("")

    add_para(doc, "What this means:", bold=True)
    add_bullet(doc, "Open files directly in Google Drive (click to view)")
    add_bullet(doc, "When you need to copy text (e.g. project instructions), open the file in Drive and copy from there")
    add_bullet(doc, "Do not download the entire folder to your computer")
    add_bullet(doc, "Do not make copies of files into your own Drive (use the shortcut instead)")
    doc.add_paragraph("")

    # ── How to use ──
    add_heading(doc, "Step 3: How to Use the Pack", level=2)

    add_para(doc, "1. Start with the Onboarding folder", bold=True)
    add_para(doc, "Open the 00-Onboarding folder and read the documents in order (00 through 12). This will get you set up and show you how everything works.")
    doc.add_paragraph("")

    add_para(doc, "2. Use the Master Index to find anything", bold=True)
    add_para(doc, "The Master Index (either the Google Sheet or the 00-MASTER-INDEX document) lists every file with a description. Use it to jump to any document quickly.")
    doc.add_paragraph("")

    add_para(doc, "3. Each folder has its own index", bold=True)
    add_para(doc, "Inside every folder there is a 00-FOLDER-INDEX document that lists everything in that folder and tells you where to go next.")
    doc.add_paragraph("")

    add_para(doc, "4. When you find a project you want to set up", bold=True)
    add_para(doc, "Open the project file, copy the instructions section, and paste it into your Claude Project on claude.ai. Replace the [PLACEHOLDER] fields with your own details.")
    doc.add_paragraph("")

    # ── Do not share ──
    add_heading(doc, "Important: Do Not Share This Folder", level=2)
    add_para(doc, "This resource is for your personal use only as an active member or opt-in.", bold=True)
    doc.add_paragraph("")
    add_bullet(doc, "Do not share the folder link with anyone")
    add_bullet(doc, "Do not forward access to team members without permission")
    add_bullet(doc, "Do not screenshot and share entire documents publicly")
    add_bullet(doc, "Do not copy and resell any of the templates or instructions")
    doc.add_paragraph("")
    add_para(doc, "If you want your team or VA to have access, contact us and we will arrange it properly. Unauthorised sharing will result in access being revoked.", color=RGBColor(0xCC, 0x44, 0x44))
    doc.add_paragraph("")

    # ── Quick reference ──
    add_heading(doc, "Quick Reference", level=2)

    add_para(doc, "Add shortcut to Drive:", bold=True)
    add_para(doc, "Folder name dropdown arrow > Organise > Add shortcut > My Drive > Add")
    doc.add_paragraph("")

    add_para(doc, "Where to start:", bold=True)
    add_para(doc, "00-Onboarding folder > read documents 00 through 12 in order")
    doc.add_paragraph("")

    add_para(doc, "Find any document:", bold=True)
    add_para(doc, "Open the Master Index (Google Sheet or 00-MASTER-INDEX.docx)")
    doc.add_paragraph("")

    add_para(doc, "Need help:", bold=True)
    add_para(doc, "Check 00-Onboarding/11-TROUBLESHOOTING-AND-FAQS for quick fixes")

    # Save
    os.makedirs(BASE_DIR, exist_ok=True)
    filepath = os.path.join(BASE_DIR, "00-READ-THIS-FIRST.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == "__main__":
    main()
