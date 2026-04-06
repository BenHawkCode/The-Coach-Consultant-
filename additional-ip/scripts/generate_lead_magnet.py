"""
Generate Lead Magnet DOCX — "The AI Business System That Replaces a £62,450/Month Team"
Branded with logo header + copyright footer, Poppins font, matching client pack styling.
Designed to blow minds and drive sign-ups WITHOUT giving away IP or instructions.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "outputs")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = (
    "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. "
    "Unauthorised sharing, copying, reproduction, or sale is prohibited "
    "without written permission. Legal action will be taken for infringements. "
    "Personal use is allowed for active members or opt-ins only."
)
FONT_NAME = "Poppins"

# Brand colours
DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
BODY_GREY = RGBColor(0x33, 0x33, 0x33)
BLUE_ACCENT = RGBColor(0x1A, 0x5A, 0xB8)
RED_ACCENT = RGBColor(0xCC, 0x00, 0x00)
GREEN_ACCENT = RGBColor(0x00, 0x66, 0x33)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
MUTED_GREY = RGBColor(0x55, 0x55, 0x55)


# ── Styling helpers (matching client pack) ──

def style_doc(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = BODY_GREY

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = DARK_NAVY
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for sn in ["List Bullet", "List Bullet 2", "List Bullet 3"]:
        if sn in doc.styles:
            doc.styles[sn].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = FONT_NAME


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = BLUE_ACCENT
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = FONT_NAME
        run_b.font.color.rgb = BLUE_ACCENT
        run_n = p.add_run(text)
        run_n.font.name = FONT_NAME
    else:
        run = p.add_run(text)
        run.font.name = FONT_NAME
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    return table


def add_spacer(doc):
    doc.add_paragraph("")


def add_divider_text(doc, text):
    """Centered muted divider text."""
    add_para(doc, text, size=9, color=MUTED_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)


# ── MAIN CONTENT ──

def main():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)

    # Set narrower margins for clean look
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ═══════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════

    add_spacer(doc)
    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "THE AI BUSINESS SYSTEM",
        bold=True, size=28, color=DARK_NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "That Replaces a \u00a362,450/Month Team",
        bold=True, size=22, color=RED_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_spacer(doc)

    add_para(
        doc,
        "How Business Owners and Service Providers Are Using Claude AI Projects "
        "to Run 34 Specialist Roles From a Single App",
        size=13, color=MUTED_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "A Copy and Paste SOP System",
        bold=True, size=14, color=DARK_NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Built for any skill level. No tech experience required. No coding. No confusion.",
        size=11, color=MUTED_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_spacer(doc)
    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "By Ben Hawksworth",
        size=12, color=DARK_NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Founder, The Coach Consultant",
        size=11, color=MUTED_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "www.thecoachconsultant.uk",
        bold=True, size=11, color=BLUE_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ═══════════════════════════════════════════
    # PAGE 2: THE PROBLEM
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "Sound Familiar?")

    add_para(doc, "You're brilliant at what you do.")
    add_spacer(doc)
    add_para(doc, "Your clients get incredible results. People ask you for advice privately all the time. You're the person everyone comes to.")
    add_spacer(doc)
    add_para(doc, "But nobody outside your existing client base actually knows you exist.")
    add_spacer(doc)
    add_para(doc, "You're working from the kitchen table. Starting at half six in the morning. Finishing at eleven at night. And 60 to 70 percent of your day is spent on operational stuff that has nothing to do with the actual transformation work you're supposed to be doing.")
    add_spacer(doc)
    add_para(doc, "The content isn't consistent. The marketing feels like an afterthought. The sales copy gets done at the last minute. Client onboarding is a mess. You know you should have SOPs but who has the time to write them when you're doing everything yourself.")
    add_spacer(doc)
    add_para(doc, "And the worst part?")
    add_spacer(doc)
    add_para(doc, "You're watching people who are genuinely less skilled than you build entire empires. Because they've got the systems sorted and you haven't.", bold=True)
    add_spacer(doc)
    add_para(doc, "You've probably tried a few AI tools already. Maybe ChatGPT. Maybe a couple of automation platforms. But nothing sticks because none of it actually sounds like you, none of it connects together properly, and half the output reads like it was written by a robot.")
    add_spacer(doc)
    add_para(doc, "I see this constantly.")
    add_spacer(doc)
    add_para(doc, "Business owners and service providers earning \u00a350K to \u00a3150K who should comfortably be at \u00a3200K plus. The results are there. The expertise is there. The market demand is there. But the systems aren't.")
    add_spacer(doc)
    add_para(doc, "That's what this document is about.")
    add_spacer(doc)
    add_para(doc, "I'm gonna walk you through exactly what I've built for my clients. A complete AI business system using Claude Projects that covers every single area of your business. Content. Marketing. Sales. Client fulfilment. Operations. Strategy. Admin. The lot.", bold=True)
    add_spacer(doc)
    add_para(doc, "Not theory. Not a list of tools to go and figure out. An actual done for you system with 34 specialist AI projects, pre-built instructions, knowledge file frameworks, and copy and paste SOPs that any person at any skill level can get set up and running.")

    # ═══════════════════════════════════════════
    # PAGE 3: WHAT IS THIS SYSTEM
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What Is This System Actually?")

    add_para(doc, "Right so let me break this down simply.")
    add_spacer(doc)
    add_para(doc, "Claude is an AI platform made by Anthropic. If you've used ChatGPT, it's a similar concept but Claude has a feature called Projects that completely changes the game for business owners.")
    add_spacer(doc)
    add_para(doc, "A Claude Project is basically a dedicated workspace where you can:")
    add_spacer(doc)
    add_bullet(doc, "Give Claude permanent instructions on how to behave, what tone to use, what to include and what to avoid", bold_prefix="Custom Instructions: ")
    add_bullet(doc, "Upload your brand guidelines, voice samples, frameworks, templates, client data, anything that Claude needs to reference every time you use it", bold_prefix="Knowledge Files: ")
    add_bullet(doc, "Claude remembers everything. No re-explaining. No repeating yourself. No starting from scratch every single conversation", bold_prefix="Persistent Context: ")
    add_spacer(doc)
    add_para(doc, "Think of it like hiring a new team member. You brief them once. Give them all the documents they need. Explain how you want things done. And from that point they just get on with it.")
    add_spacer(doc)
    add_para(doc, "That's what Claude Projects does. Except you're not hiring one team member.")
    add_spacer(doc)
    add_para(doc, "You're hiring 34.", bold=True, size=14)
    add_spacer(doc)
    add_para(doc, "Each one is a specialist. Each one knows your brand inside out. Each one follows your exact voice, tone and style. Each one works 24/7. No sick days. No holiday cover. No management overhead.")
    add_spacer(doc)
    add_para(doc, "And the entire system costs \u00a315 a month.", bold=True, size=13, color=GREEN_ACCENT)

    # ═══════════════════════════════════════════
    # PAGE 4: THE 34 PROJECTS OVERVIEW
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What You Actually Get: 34 Specialist AI Projects")

    add_para(doc, "This system is split across 8 categories covering every area of your business. Each project is a fully built AI specialist that handles a specific function.")
    add_spacer(doc)
    add_para(doc, "I'm not gonna give you the instructions or the knowledge files here. That's the IP. But I am gonna show you exactly what each one does so you can see the full picture of what this system handles for you.", size=11, color=MUTED_GREY)

    # ── Category 1: Content Marketing ──
    add_spacer(doc)
    add_heading(doc, "Category 1: Content Marketing", level=2)
    add_para(doc, "Replaces: Content Strategist, Copywriter, Email Specialist, Lead Gen Expert, Community Manager", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a311,400/month (\u00a3136,800/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["1. Content Strategy & Calendar", "Plans your monthly themes, maps content pillars, generates weekly calendars with platform specific posts already assigned"],
            ["2. Blog & Long Form Content", "Writes SEO optimised blog posts, articles, guides and thought leadership pieces in your exact voice"],
            ["3. Email Marketing Sequences", "Drafts complete welcome sequences, nurture campaigns, launch emails, newsletters and re-engagement flows"],
            ["4. Lead Magnets & Opt-In Content", "Creates checklists, guides, templates, quizzes and mini courses designed to attract your ideal clients"],
            ["33. Community Management", "Generates group rules, welcome sequences, engagement prompts, moderation guidelines and member nurture content"],
        ],
    )

    # ── Category 2: Social Media ──
    add_spacer(doc)
    add_heading(doc, "Category 2: Social Media", level=2)
    add_para(doc, "Replaces: Social Media Manager, LinkedIn Strategist, YouTube Producer, Content Repurposer, Podcast Producer", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a39,700/month (\u00a3116,400/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["5. Instagram Content Creator", "Generates carousel ideas, reel scripts, captions, stories, hashtag strategies and engagement driving CTAs"],
            ["6. LinkedIn Thought Leadership", "Writes posts, articles, comment responses, connection messages and newsletter content that positions you as the authority"],
            ["7. YouTube Content & Scripts", "Plans video topics, writes full teleprompter scripts with timestamps, creates thumbnail titles and descriptions"],
            ["8. Content Repurposing Engine", "Takes one piece of content and turns it into platform specific versions for every channel you use"],
            ["32. Podcast Strategy & Appearances", "Develops episode topics, guest bios, interview questions, pitch templates and show notes"],
        ],
    )

    # ── Category 3: Sales & Lead Generation ──
    add_spacer(doc)
    add_heading(doc, "Category 3: Sales & Lead Generation", level=2)
    add_para(doc, "Replaces: Sales Copywriter, Proposal Writer, Outreach Specialist, Webinar Producer", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a38,300/month (\u00a399,600/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["9. Sales Landing Page Copy", "Writes high converting sales pages, landing pages, webinar registration pages and checkout copy"],
            ["10. Proposal & Pitch Writer", "Drafts client proposals, pitch decks, discovery follow ups and scope of work documents"],
            ["11. DM & Outreach Sequences", "Crafts personalised outreach messages, follow up sequences and networking scripts for Instagram, LinkedIn and email"],
            ["12. Webinar & Workshop Planner", "Plans content, writes scripts, creates slide outlines, follow up emails and replay sequences"],
        ],
    )

    # ── Category 4: Client Fulfilment ──
    doc.add_page_break()
    add_heading(doc, "Category 4: Client Fulfilment & Delivery", level=2)
    add_para(doc, "Replaces: Client Success Coordinator, Session Admin, Reporting Analyst, Course Creator, Retention Specialist", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a38,800/month (\u00a3105,600/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["13. Client Onboarding Assistant", "Generates welcome packets, onboarding emails, intake forms, kick off agendas and expectation setting documents"],
            ["14. Coaching Session Prep & Notes", "Prepares session agendas, summarises progress, drafts action items and tracks milestones"],
            ["15. Client Reporting & Results", "Builds monthly reports, progress summaries, ROI breakdowns and strategy recommendations"],
            ["16. Course & Programme Content", "Develops module outlines, lesson scripts, workbook content, quizzes and community prompts"],
            ["31. Client Offboarding & Retention", "Handles offboarding sequences, feedback collection, referral systems and win back campaigns"],
        ],
    )

    # ── Category 5: Operations ──
    add_spacer(doc)
    add_heading(doc, "Category 5: Operations & Business Management", level=2)
    add_para(doc, "Replaces: Operations Manager, Business Analyst, HR Coordinator, Bookkeeper, Financial Controller, Automation Specialist", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a313,500/month (\u00a3162,000/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["17. SOP & Process Documentation", "Creates standard operating procedures, workflow documentation, training guides and delegation checklists"],
            ["18. KPI & Business Reporting", "Analyses metrics, generates performance reports, spots trends and drafts strategic recommendations"],
            ["19. Team Communication & Management", "Drafts team updates, meeting agendas, feedback reviews, hiring briefs and contractor briefs"],
            ["20. Finance & Invoicing Assistant", "Handles invoice follow ups, expense categorisation, financial summaries and budget planning"],
            ["29. Finance & Strategic Reporting", "Produces financial analysis, forecasting, cash flow projections and strategic finance reviews"],
            ["34. Tech Stack & Automation Setup", "Documents your tech stack, maps integrations, plans automation workflows and troubleshoots systems"],
        ],
    )

    # ── Category 6: Strategy & Growth ──
    add_spacer(doc)
    add_heading(doc, "Category 6: Strategy & Growth", level=2)
    add_para(doc, "Replaces: Business Consultant, Brand Strategist, Business Dev Manager, Research Analyst, Pricing Consultant", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a312,600/month (\u00a3151,200/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["21. Business Strategy & Planning", "Brainstorms offers, maps revenue streams, plans quarterly goals, conducts market analysis"],
            ["22. Brand Positioning Strategist", "Refines messaging, unique value proposition, brand story and competitive positioning"],
            ["23. Partnership & Collaboration", "Drafts partnership proposals, affiliate details, guest podcast pitches and joint venture outreach"],
            ["24. Customer Research & Insights", "Analyses surveys, extracts testimonial themes, identifies pain points, generates market insight reports"],
            ["30. Offer Design & Pricing Strategy", "Structures offers, models pricing, builds value stacks and creates comparison frameworks"],
        ],
    )

    # ── Category 7: Admin & Organisation ──
    doc.add_page_break()
    add_heading(doc, "Category 7: Admin & Organisation", level=2)
    add_para(doc, "Replaces: Communications Manager, Event Coordinator, Legal Administrator, Virtual Assistant", size=10, color=MUTED_GREY)
    add_para(doc, "Staff cost replaced: \u00a36,200/month (\u00a374,400/year)", bold=True)
    add_spacer(doc)

    add_table(doc,
        ["Project", "What It Does For You"],
        [
            ["25. Email & Communication Manager", "Drafts client emails, responds to enquiries, writes follow ups and manages correspondence"],
            ["26. Meeting & Event Coordinator", "Creates agendas, event run sheets, speaker briefs, post event summaries and attendee follow ups"],
            ["27. Legal & Compliance Drafts", "Drafts terms of service, privacy policies, contract clauses, refund policies and disclaimer text"],
            ["28. Personal Productivity & Planning", "Plans your week, prioritises tasks, drafts daily action plans, sets quarterly goals"],
        ],
    )

    # ── Category 8: Claude Code Guides ──
    add_spacer(doc)
    add_heading(doc, "Category 8: Maintenance & Update Guides", level=2)
    add_para(doc, "12 step by step guides that keep your entire system current without needing any technical knowledge.", size=11, color=MUTED_GREY)
    add_spacer(doc)

    add_table(doc,
        ["Guide", "What It Covers"],
        [
            ["Full Rebrand Update", "Change your entire brand across all 34 projects in one go"],
            ["Bulk CTA Update", "Update every call to action across the system when your offer changes"],
            ["Pricing & Offer Update", "Roll out new pricing or restructured offers system wide"],
            ["Target Audience Update", "Shift your audience targeting across all content projects"],
            ["Quarterly Health Check", "Review and optimise your entire system every 90 days"],
            ["New Knowledge File Rollout", "Add new brand documents, transcripts or frameworks across projects"],
            ["Voice & Tone Recalibration", "Fine tune your AI voice when your natural style evolves"],
            ["Seasonal Campaign Setup", "Spin up campaign specific projects for launches and promotions"],
            ["Project Backup & Export", "Back up your entire system so nothing gets lost"],
            ["New Project Setup", "Add brand new projects following the exact same framework"],
            ["Compliance & Disclaimer Update", "Update legal and compliance language across all projects"],
            ["Team & VA Handover Pack", "Hand the entire system to a VA or team member with zero confusion"],
        ],
    )

    # ═══════════════════════════════════════════
    # THE NUMBERS
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "The Numbers That Matter")

    add_para(doc, "Right so let me spell this out properly.")
    add_spacer(doc)
    add_para(doc, "If you were to hire actual specialists to do what these 34 projects do, here's what you'd be paying. Not entry level hires. Not interns. Experienced professionals who actually know what they're doing.")
    add_spacer(doc)

    add_para(
        doc,
        "TOTAL MONTHLY STAFF COST THIS SYSTEM REPLACES:",
        bold=True, size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "\u00a362,450/month",
        bold=True, size=28, color=RED_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "That's \u00a3749,400 per year in staff wages alone.",
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_spacer(doc)

    add_para(doc, "And that's before you factor in employer National Insurance contributions at 13.8 percent, pension auto-enrolment at 3 to 5 percent, recruitment fees, software licenses, desk space, management time, sick days and holiday cover.")
    add_spacer(doc)
    add_para(doc, "With those hidden costs added, the real figure is closer to \u00a3870,000 to \u00a3910,000 per year.", bold=True)
    add_spacer(doc)

    add_heading(doc, "Your Actual Cost to Run This Entire System", level=2)
    add_spacer(doc)

    add_table(doc,
        ["Tool", "What It Does", "Monthly Cost"],
        [
            ["Claude Pro", "Powers all 34 AI projects", "\u00a315/month"],
            ["Make.com (Pro Plan)", "Connects your tools and runs automations", "\u00a316/month"],
            ["Total", "", "\u00a331/month"],
        ],
    )
    add_spacer(doc)

    add_para(
        doc,
        "\u00a331/month replaces \u00a362,450/month in staff wages.",
        bold=True, size=13, color=GREEN_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "That is a 99.95% cost reduction.",
        bold=True, size=14, color=GREEN_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_spacer(doc)

    # Side by side comparison
    add_heading(doc, "The Comparison", level=2)
    add_spacer(doc)

    add_table(doc,
        ["", "Hiring Staff", "This System"],
        [
            ["Monthly Cost", "\u00a362,450", "\u00a331"],
            ["Annual Cost", "\u00a3749,400", "\u00a3372"],
            ["Hidden Costs (NI, pension, recruitment)", "\u00a3120,000 to \u00a3160,000", "\u00a30"],
            ["True Annual Cost", "\u00a3870,000 to \u00a3910,000", "\u00a3372"],
            ["Availability", "Business hours minus holidays and sick days", "24/7, 365 days"],
            ["Ramp Up Time", "3 to 6 months per hire", "Same day"],
            ["Management Overhead", "5 to 10 hours per week", "Zero"],
            ["Recruitment Cost", "\u00a33,000 to \u00a38,000 per role", "Zero"],
            ["Risk of Bad Hire", "Real and expensive", "None"],
        ],
    )

    # ═══════════════════════════════════════════
    # WHAT BUSINESS OWNERS ACTUALLY SPEND
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What Most Business Owners Actually Spend Right Now")

    add_para(doc, "Nobody hires 34 specialists. Obviously. That's the whole point.")
    add_spacer(doc)
    add_para(doc, "What actually happens is you hire one or two people, outsource a couple of things to freelancers, and do everything else yourself. Which means most of your business isn't covered properly and you're working until eleven at night trying to fill the gaps.")
    add_spacer(doc)
    add_para(doc, "I see this all the time with business owners at the \u00a350K to \u00a3150K level. The typical monthly spend looks something like this:")
    add_spacer(doc)

    add_table(doc,
        ["Typical Hire", "Monthly Cost"],
        [
            ["Social Media Manager", "\u00a31,600"],
            ["VA or Admin Support", "\u00a31,100"],
            ["Freelance Copywriter", "\u00a32,500"],
            ["Typical Monthly Spend", "\u00a35,200"],
        ],
    )
    add_spacer(doc)

    add_para(doc, "That \u00a35,200 per month gets you 3 out of 34 roles covered. Partially.")
    add_spacer(doc)
    add_para(doc, "This system gives you all 34 for \u00a331 a month.", bold=True, size=13)
    add_spacer(doc)
    add_para(doc, "And every single one of them knows your brand, follows your voice, understands your audience, and works to the exact same standard every time you use it.")

    # ═══════════════════════════════════════════
    # WHY THIS IS DIFFERENT
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "Why This Is Different to Anything Else Out There")

    add_para(doc, "I know what you're thinking. There are thousands of AI tools, courses and templates floating around. What makes this any different.")
    add_spacer(doc)
    add_para(doc, "The problem is this.", bold=True)
    add_spacer(doc)
    add_para(doc, "Most AI solutions give you a tool and leave you to figure it out. Or they give you a list of prompts that sound nothing like you and produce generic output that could have come from anyone's business.")
    add_spacer(doc)
    add_para(doc, "This system is different because it's not a tool. It's not a course. It's not a list of prompts.")
    add_spacer(doc)
    add_para(doc, "It's a complete, pre-built business operating system.", bold=True, size=13)
    add_spacer(doc)

    add_heading(doc, "What makes it different:", level=2)
    add_spacer(doc)

    add_bullet(doc, " Every project comes with pre-written instructions that tell Claude exactly how to behave for that specific function. You don't write anything from scratch. You copy, paste, customise the placeholders with your business details, and you're running.", bold_prefix="Copy and paste setup.")
    add_spacer(doc)
    add_bullet(doc, " The system is designed so that someone who has never used AI before can get this up and running. Every step is documented. Every instruction is written for you. Every knowledge file is mapped out. No coding. No tech experience. No guesswork.", bold_prefix="Built for any skill level.")
    add_spacer(doc)
    add_bullet(doc, " This isn't generic AI that sounds like a robot. Every project is built around your voice, your brand, your audience and your specific business. The output sounds like you wrote it.", bold_prefix="Your voice, your brand, your business.")
    add_spacer(doc)
    add_bullet(doc, " 34 projects covering every single area of your business. Content marketing. Social media. Sales. Client fulfilment. Operations. Strategy. Admin. Plus 12 maintenance guides to keep everything current. Nothing is left uncovered.", bold_prefix="Complete coverage.")
    add_spacer(doc)
    add_bullet(doc, " This isn't a one off template that goes stale. The system grows with your business. The maintenance guides show you exactly how to update everything as your brand evolves, your offers change, or your audience shifts.", bold_prefix="A living system that evolves.")
    add_spacer(doc)
    add_bullet(doc, " The onboarding pack walks you through setup in phases. You start with the highest impact projects and build from there. You'll be getting value from this system within your first 15 minutes.", bold_prefix="Structured onboarding.")

    # ═══════════════════════════════════════════
    # THE ONBOARDING SYSTEM
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "The Onboarding System")

    add_para(doc, "One of the biggest concerns people have is whether they'll actually be able to set this up. Especially if they're not particularly technical.")
    add_spacer(doc)
    add_para(doc, "That's exactly why the onboarding system exists.", bold=True)
    add_spacer(doc)
    add_para(doc, "The client pack includes a complete onboarding folder with 13 documents that take you from zero to fully operational. Everything is step by step. Everything is copy and paste. Nothing requires any technical skill whatsoever.")
    add_spacer(doc)

    add_heading(doc, "What's in the onboarding pack:", level=2)
    add_spacer(doc)

    add_bullet(doc, " The full cost savings analysis you've seen in this document, personalised to show what you're specifically replacing", bold_prefix="Cost Savings Breakdown.")
    add_spacer(doc)
    add_bullet(doc, " Get your first project live and producing output within 15 minutes of opening the pack", bold_prefix="Your First 15 Minutes Guide.")
    add_spacer(doc)
    add_bullet(doc, " Complete orientation with screenshots showing you exactly where to click, what to type, and what to upload", bold_prefix="Start Here Guide.")
    add_spacer(doc)
    add_bullet(doc, " A structured worksheet that captures your brand voice, messaging, audience and positioning before you set up a single project. This is what makes the output sound like you instead of a robot", bold_prefix="Brand Definition Worksheet.")
    add_spacer(doc)
    add_bullet(doc, " Honest breakdown of which Claude plan you need and why, so you're not paying for features you don't use", bold_prefix="Why You Need Pro Explained.")
    add_spacer(doc)
    add_bullet(doc, " Exactly which additional tools and browser setup will get you the best results from day one", bold_prefix="Recommended Tools & Setup.")
    add_spacer(doc)
    add_bullet(doc, " How to organise and upload the documents that make Claude actually understand your business", bold_prefix="Knowledge File Starter Kit.")
    add_spacer(doc)
    add_bullet(doc, " A day by day setup sequence so you're not overwhelmed. Day one focuses on brand setup. By day five you've got your highest impact projects live", bold_prefix="Your First Week Framework.")
    add_spacer(doc)
    add_bullet(doc, " Five rules for writing prompts that get the best possible output every time", bold_prefix="How to Get Best Results.")
    add_spacer(doc)
    add_bullet(doc, " A single page overview of all 34 projects with ROI ratings so you know which ones to set up first", bold_prefix="All 34 Projects at a Glance.")
    add_spacer(doc)
    add_bullet(doc, " A 5 phase implementation checklist that tracks your progress from setup through to full optimisation", bold_prefix="Setup Progress Tracker.")
    add_spacer(doc)
    add_bullet(doc, " Solutions to the most common questions and issues people run into during setup", bold_prefix="Troubleshooting & FAQs.")
    add_spacer(doc)
    add_bullet(doc, " A per project checklist to verify everything is configured properly before you start using it", bold_prefix="Project Review Checklist.")

    # ═══════════════════════════════════════════
    # WHAT YOUR WEEK LOOKS LIKE
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What Your Week Actually Looks Like With This System")

    add_para(doc, "The shift is simple.")
    add_spacer(doc)
    add_para(doc, "Right now you're spending 60 to 70 percent of your time on operational work. Content creation. Email writing. Admin. Reporting. All the stuff that keeps the lights on but doesn't actually move the business forward.")
    add_spacer(doc)
    add_para(doc, "With this system running, that flips.", bold=True)
    add_spacer(doc)

    add_heading(doc, "Monday:", level=3)
    add_para(doc, "You open your Content Strategy project. Ask Claude to plan your week. It already knows your content pillars, your audience, your posting schedule and your brand voice. Within minutes you've got a full week of content mapped across every platform. Captions written. Carousel copy done. Email drafted. YouTube video outlined.")
    add_spacer(doc)

    add_heading(doc, "Tuesday:", level=3)
    add_para(doc, "Client session in the morning. Afterwards you open your Session Prep project. Claude summarises the session, drafts follow up action items, updates the client's progress notes. What used to take 45 minutes takes 5.")
    add_spacer(doc)

    add_heading(doc, "Wednesday:", level=3)
    add_para(doc, "New lead comes in. You open your Proposal Writer project. Give it the discovery call notes. It drafts a complete, branded proposal in your voice. You review it, send it. Done before lunch.")
    add_spacer(doc)

    add_heading(doc, "Thursday:", level=3)
    add_para(doc, "Quarterly review day. You open your KPI Reporting project. Feed in your numbers. It generates a full performance report with trends, insights and recommendations. Then you open Business Strategy and plan next quarter based on actual data.")
    add_spacer(doc)

    add_heading(doc, "Friday:", level=3)
    add_para(doc, "Batch your social media for next week. Open Instagram Creator, LinkedIn Thought Leadership and Content Repurposing. Three projects. One hour. An entire week of platform specific content ready to schedule.")
    add_spacer(doc)

    add_para(doc, "That's how you go from working 60 hours a week to working 25 to 30. Not by doing less. By having a system that handles the 60 to 70 percent of work that was eating your time.", bold=True)

    # ═══════════════════════════════════════════
    # WHAT THIS DOESN'T DO
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What This System Doesn't Do")

    add_para(doc, "I want to be completely straight with you here because I'm not interested in overselling this.")
    add_spacer(doc)

    add_para(doc, "This system does not replace:", bold=True)
    add_spacer(doc)
    add_bullet(doc, "You still make the strategic calls. Claude prepares the analysis, drafts the options, organises the data. But the decisions are yours")
    add_bullet(doc, "If you're a coach or consultant, your clients are paying for you. The system handles everything around the delivery but the actual transformation work is still human")
    add_bullet(doc, "AI doesn't build relationships. It handles the admin around them. The follow ups, the onboarding, the reporting. But the real connection is still you")
    add_bullet(doc, "Everything Claude produces should be reviewed before it goes out. The quality is consistently high but a human eye catches nuance that AI sometimes misses")
    add_spacer(doc)

    add_para(doc, "What it DOES replace:", bold=True)
    add_spacer(doc)
    add_bullet(doc, "The first draft of literally everything your business produces")
    add_bullet(doc, "Research, planning and preparation across every area of your business")
    add_bullet(doc, "Repetitive admin and operational tasks that eat your evenings")
    add_bullet(doc, "Content creation across every single channel you use")
    add_bullet(doc, "Documentation, reporting and analysis that you never have time for")
    add_bullet(doc, "The 60 to 70 percent of your week currently spent on work that should be systematised")
    add_spacer(doc)

    add_para(doc, "The bottom line is this. You're still the expert. You're still the person your clients trust. But instead of spending your entire week on operational work, you spend your time on the stuff that actually matters. The transformation work, the relationship building, the strategic thinking that grows your business.", bold=True)

    # ═══════════════════════════════════════════
    # WHO THIS IS FOR
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "Who This Is Built For")

    add_para(doc, "This system was specifically built for business owners and service providers who:")
    add_spacer(doc)

    add_bullet(doc, "Are earning \u00a350K to \u00a3150K and know they should be at \u00a3200K plus based on the quality of work they deliver")
    add_bullet(doc, "Are working from the kitchen table, doing everything themselves, starting early and finishing late")
    add_bullet(doc, "Have tried AI tools before but nothing stuck because it all sounded generic or couldn't be integrated properly")
    add_bullet(doc, "Know they need systems but don't have the time or technical knowledge to build them from scratch")
    add_bullet(doc, "Are watching competitors with less skill and worse results build bigger businesses because they've got better systems")
    add_bullet(doc, "Want a complete solution, not another tool to figure out")
    add_spacer(doc)

    add_para(doc, "If you're a coach doing \u00a3500K plus with a full team already in place, this probably isn't for you. You've already got the infrastructure.")
    add_spacer(doc)
    add_para(doc, "But if you're the person who's brilliant at what you do, your clients get incredible results, and you're stuck doing everything manually because you haven't found the right system yet?")
    add_spacer(doc)
    add_para(doc, "This is exactly what you've been looking for.", bold=True, size=13)

    # ═══════════════════════════════════════════
    # NEXT STEPS
    # ═══════════════════════════════════════════

    doc.add_page_break()

    add_heading(doc, "What Happens Next")

    add_para(doc, "You've just seen the full picture of what this system does.")
    add_spacer(doc)
    add_para(doc, "34 specialist AI projects. 8 business categories. 12 maintenance guides. A complete onboarding system. Copy and paste SOPs. Built for any skill level.")
    add_spacer(doc)
    add_para(doc, "All replacing what would cost \u00a362,450 a month in staff wages for \u00a331 a month in software.", bold=True)
    add_spacer(doc)
    add_para(doc, "The question isn't whether AI can help your business. You already know it can.")
    add_spacer(doc)
    add_para(doc, "The question is whether you want to spend the next six months trying to figure it out yourself, watching competitors pull further ahead whilst you debate which tool to use.")
    add_spacer(doc)
    add_para(doc, "Or whether you want someone who's already built this, tested this, and deployed this for hundreds of business owners to set it up for your business properly.")
    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "Book a call and I'll walk you through exactly how this would work for your specific business.",
        bold=True, size=13, color=DARK_NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_spacer(doc)
    add_para(
        doc,
        "www.thecoachconsultant.uk",
        bold=True, size=16, color=BLUE_ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_spacer(doc)
    add_spacer(doc)

    add_divider_text(doc, "No pressure. No hard sell. Just a conversation about what this system could do for you specifically.")

    add_spacer(doc)
    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "Ben Hawksworth",
        bold=True, size=12, color=DARK_NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Founder, The Coach Consultant",
        size=11, color=MUTED_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_spacer(doc)
    add_spacer(doc)

    add_para(
        doc,
        "All salary data based on 2025-2026 UK specialist and freelancer market rates for experienced "
        "professionals with 3 to 5 plus years experience. Costs reflect fractional or part time "
        "engagement typical for small businesses and solo operators. Sources include Robert Half UK "
        "Salary Guide, Glassdoor UK, Reed.co.uk and YunoJuno Freelancer Rates Report.",
        size=8, color=LIGHT_GREY,
    )

    # ── Save ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filepath = os.path.join(OUTPUT_DIR, "LEAD-MAGNET-AI-BUSINESS-SYSTEM.docx")
    doc.save(filepath)
    print(f"\nCreated: {filepath}")
    print(f"Pages: ~15 (branded DOCX ready for PDF export)")


if __name__ == "__main__":
    main()
