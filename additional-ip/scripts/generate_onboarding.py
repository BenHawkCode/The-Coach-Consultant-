"""
Generate Onboarding Documents for Claude Projects Client Pack.
Creates the 00-Onboarding folder with all documents new members need
to get started, upgrade their plan, and maximise their experience.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "Claude-Projects-Client-Pack")
ONBOARDING_DIR = os.path.join(BASE_DIR, "00-Onboarding")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "logo-white.png")
FOOTER_TEXT = "\u00a9 The Coach Consultant FitPro Ltd. All rights reserved. Unauthorised sharing, copying, reproduction, or sale is prohibited without written permission. Legal action will be taken for infringements. Personal use is allowed for active members or opt-ins only."
FONT_NAME = "Poppins"


# ─── Styling helpers (matched to existing pack styling) ───

def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            h.font.size = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
        else:
            h.font.size = Pt(13)

    for style_name in ['List Bullet', 'List Bullet 2', 'List Bullet 3']:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = FONT_NAME


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))


def add_footer_copyright(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = FONT_NAME


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if bold:
        run.bold = True
        if not color:
            run.font.color.rgb = RGBColor(0x1a, 0x5a, 0xb8)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = FONT_NAME
    return p


def add_checkbox(doc, text):
    doc.add_paragraph(f"\u2610  {text}", style='List Bullet')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    return table


def save_doc(doc, filename):
    os.makedirs(ONBOARDING_DIR, exist_ok=True)
    filepath = os.path.join(ONBOARDING_DIR, filename)
    doc.save(filepath)
    print(f"  Created: {filepath}")


def new_doc():
    doc = Document()
    style_doc(doc)
    add_header_logo(doc)
    add_footer_copyright(doc)
    return doc


# ─── Document 1: Why You Need Claude Pro (& Max) ───

def create_plan_guide():
    doc = new_doc()

    add_heading(doc, "Why You Need Claude Pro (and When to Go Max)")
    add_para(doc, "Your guide to picking the right Claude plan so this entire system actually works for you.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # The short version
    add_heading(doc, "The Short Version", level=2)
    add_para(doc, "You need Claude Pro at minimum.", bold=True, size=13)
    add_para(doc, "The free plan will not cut it. You will run out of messages within the first hour of setting up your projects and hit limitations that make this entire system frustrating instead of powerful.")
    doc.add_paragraph("")
    add_para(doc, "If you are serious about running your business on this system, Max is the plan that removes all the friction.")
    doc.add_paragraph("")

    # What each plan gives you
    add_heading(doc, "What Each Plan Actually Gives You", level=2)
    doc.add_paragraph("")

    add_table(doc,
        ['Feature', 'Free Plan', 'Pro Plan (\u00a315/month)', 'Max Plan (\u00a345/month)'],
        [
            ['Messages per day', 'Very limited (runs out fast)', '5x more than free', '20x more than free'],
            ['Claude Projects', 'No access', 'Full access', 'Full access'],
            ['Knowledge file uploads', 'No', 'Yes (up to 200MB per project)', 'Yes (up to 200MB per project)'],
            ['Custom instructions', 'No', 'Yes', 'Yes'],
            ['Model access', 'Claude Sonnet only', 'Claude Sonnet + Opus', 'Claude Sonnet + Opus (priority)'],
            ['Priority access', 'No (queued during busy times)', 'Yes', 'Always available'],
            ['Extended thinking', 'No', 'Limited', 'Full access'],
            ['Usage during peak hours', 'Often throttled or unavailable', 'Occasional rate limits', 'Rarely limited'],
        ]
    )
    doc.add_paragraph("")

    # Why free doesn't work
    add_heading(doc, "Why the Free Plan Does Not Work for This System", level=2)
    add_para(doc, "This system is built on Claude Projects. The free plan does not have Projects.", bold=True)
    doc.add_paragraph("")
    add_para(doc, "Without Projects you cannot:")
    add_bullet(doc, "Set custom instructions (the entire point of each project file)")
    add_bullet(doc, "Upload knowledge files (your brand voice, client data, SOPs)")
    add_bullet(doc, "Create separate workspaces for different business functions")
    add_bullet(doc, "Get Claude to remember your context between conversations")
    doc.add_paragraph("")
    add_para(doc, "You would be starting from scratch every single conversation. Explaining who you are, what your business does, how you speak, and what you need. Every. Single. Time.")
    doc.add_paragraph("")
    add_para(doc, "That is not a system. That is a headache.")
    doc.add_paragraph("")

    # Why Pro is the minimum
    add_heading(doc, "Why Pro Is the Minimum", level=2)
    add_para(doc, "Claude Pro gives you everything you need to run this system:")
    doc.add_paragraph("")
    add_bullet(doc, "Full access to Claude Projects with custom instructions and knowledge files")
    add_bullet(doc, "Enough daily messages to set up and use multiple projects")
    add_bullet(doc, "Access to Claude Opus (the most capable model) for complex tasks")
    add_bullet(doc, "Priority access so you are not waiting in queues during peak times")
    doc.add_paragraph("")
    add_para(doc, "For \u00a315/month you are getting a system that replaces \u00a362,450/month in specialist roles. That is not a cost. That is the best return on investment in your entire business.", bold=True)
    doc.add_paragraph("")

    # When to upgrade to Max
    doc.add_page_break()
    add_heading(doc, "When to Upgrade to Max", level=2)
    add_para(doc, "Start with Pro. Upgrade to Max when you notice any of these:", bold=True)
    doc.add_paragraph("")
    add_bullet(doc, "You are hitting message limits regularly (running out of messages mid-afternoon)")
    add_bullet(doc, "You are using Claude for content creation AND operations AND strategy daily")
    add_bullet(doc, "You are running 10+ active projects and switching between them throughout the day")
    add_bullet(doc, "You want extended thinking for complex strategy and planning work")
    add_bullet(doc, "You are building this into your daily workflow and Claude is genuinely your first port of call")
    doc.add_paragraph("")
    add_para(doc, "Most people who take this system seriously end up on Max within 2-4 weeks. Not because they have to, but because once you see what Claude can do with proper setup, you want more of it.")
    doc.add_paragraph("")

    # The real cost comparison
    add_heading(doc, "The Real Cost Comparison", level=2)
    doc.add_paragraph("")
    add_table(doc,
        ['', 'What You Pay', 'What You Get'],
        [
            ['Free Plan', '\u00a30/month', 'No Projects, no system, starting from scratch every time'],
            ['Pro Plan', '\u00a315/month', 'Full system access, 34 AI specialists, 5x message allowance'],
            ['Max Plan', '\u00a345/month', 'Unlimited-feel usage, priority everything, extended thinking'],
            ['Hiring 1 VA', '\u00a31,100/month', 'One person covering basic admin tasks only'],
            ['Hiring 1 copywriter', '\u00a32,500/month', 'One person writing content for one channel'],
            ['Hiring 3 specialists', '\u00a35,200/month', 'Three roles out of 34 covered, partially'],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Even the Max plan at \u00a345/month is less than 1% of what a single specialist would cost you.", bold=True)
    doc.add_paragraph("")

    # How to upgrade
    add_heading(doc, "How to Upgrade Your Plan", level=2)
    add_para(doc, "Step 1: Go to claude.ai and log in")
    add_para(doc, "Step 2: Click your profile icon in the bottom left")
    add_para(doc, "Step 3: Click 'Settings'")
    add_para(doc, "Step 4: Click 'Billing' or 'Subscription'")
    add_para(doc, "Step 5: Choose Pro or Max and follow the payment steps")
    doc.add_paragraph("")
    add_para(doc, "It takes less than 2 minutes. Once upgraded, Projects appears in your left sidebar immediately.")
    doc.add_paragraph("")

    # Team plans
    add_heading(doc, "A Note on Team Plans", level=2)
    add_para(doc, "If you have a team or VA who will also be using Claude, Anthropic offers Team plans that give everyone their own Pro-level access with shared billing. Worth considering once you have the system running and want to hand projects off to team members.")

    save_doc(doc, "04-WHY-YOU-NEED-PRO-AND-MAX.docx")


# ─── Document 2: Recommended Tools & Setup ───

def create_tools_guide():
    doc = new_doc()

    add_heading(doc, "Recommended Tools and Setup")
    add_para(doc, "The tools that will make your experience with this system faster, easier, and more enjoyable.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Voice to AI
    add_heading(doc, "Voice-to-AI: Stop Typing, Start Talking", level=2)
    add_para(doc, "This is the single biggest upgrade you can make to how you use Claude.", bold=True)
    doc.add_paragraph("")

    add_para(doc, "Wispr Flow (Recommended)", bold=True)
    add_bullet(doc, "What it does: Lets you speak naturally and it types for you, anywhere on your computer")
    add_bullet(doc, "Why it matters: You can brief Claude the way you would brief a team member, by talking")
    add_bullet(doc, "How it works: Press a hotkey, speak your prompt, and Wispr Flow transcribes it directly into the Claude chat box")
    add_bullet(doc, "Cost: Free tier available, paid plans for heavier use")
    add_bullet(doc, "Get it: wispr.flow")
    doc.add_paragraph("")
    add_para(doc, "This completely changes the experience. Instead of staring at a blank text box trying to type the perfect prompt, you just talk. Tell Claude what you need the way you would tell a real person. The results are better because your prompts are more natural and detailed when you speak them.")
    doc.add_paragraph("")

    add_para(doc, "Other Voice-to-Text Options:", bold=True)
    add_bullet(doc, "macOS Dictation (built in, free): Press Fn key twice to activate. Good enough to start with")
    add_bullet(doc, "Windows Voice Typing (built in, free): Press Windows + H to activate")
    add_bullet(doc, "Otter.ai: Great for recording meetings and calls, then pasting transcripts into Claude")
    doc.add_paragraph("")

    # Browser setup
    add_heading(doc, "Browser Setup for Speed", level=2)
    add_para(doc, "Pin Claude as a tab:", bold=True)
    add_bullet(doc, "Right-click the Claude tab in your browser and select 'Pin Tab'")
    add_bullet(doc, "This keeps Claude one click away at all times without cluttering your tab bar")
    doc.add_paragraph("")
    add_para(doc, "Bookmark your most-used projects:", bold=True)
    add_bullet(doc, "Open each project in Claude, then bookmark the URL")
    add_bullet(doc, "Create a browser bookmark folder called 'Claude Projects' for quick access")
    add_bullet(doc, "Your top 3-5 projects should be one click away")
    doc.add_paragraph("")
    add_para(doc, "Use keyboard shortcuts:", bold=True)
    add_bullet(doc, "Cmd+K (Mac) or Ctrl+K (Windows) in Claude opens the project switcher")
    add_bullet(doc, "Learn this shortcut and you will fly between projects")
    doc.add_paragraph("")

    # File preparation
    add_heading(doc, "Preparing Your Knowledge Files", level=2)
    add_para(doc, "The quality of what Claude produces depends directly on what you feed it. Before you start setting up projects, prepare these files:")
    doc.add_paragraph("")
    add_para(doc, "Priority 1: Your Voice Files", bold=True)
    add_bullet(doc, "Record yourself explaining what your business does for 5-10 minutes (voice memo on your phone)")
    add_bullet(doc, "Transcribe it using Wispr Flow, Otter.ai, or even just the free voice typing on your phone")
    add_bullet(doc, "Save it as a text file or PDF and upload it to every writing-based project")
    add_bullet(doc, "This single file will transform the quality of everything Claude writes for you")
    doc.add_paragraph("")
    add_para(doc, "Priority 2: Your Best Content", bold=True)
    add_bullet(doc, "Grab your 5-10 best Instagram captions, emails, or posts")
    add_bullet(doc, "Copy them into a single document")
    add_bullet(doc, "Upload to any project that creates content for you")
    doc.add_paragraph("")
    add_para(doc, "Priority 3: Your Business Basics", bold=True)
    add_bullet(doc, "One document covering: what you sell, who you sell it to, your prices, and your website")
    add_bullet(doc, "Does not need to be polished. Bullet points are fine")
    add_bullet(doc, "Upload to every project so Claude always has your business context")
    doc.add_paragraph("")

    # Note-taking
    add_heading(doc, "Capturing Ideas on the Go", level=2)
    add_para(doc, "You will constantly think of things to ask Claude throughout the day. Have a system for capturing them:")
    doc.add_paragraph("")
    add_bullet(doc, "Apple Notes or Google Keep: Quick and simple. Create a note called 'Claude Prompts' and dump ideas there")
    add_bullet(doc, "Voice memos: Record a 30-second voice note, then transcribe it later and paste into Claude")
    add_bullet(doc, "Notion or similar: If you already use a project management tool, create a 'Claude Tasks' page")
    doc.add_paragraph("")
    add_para(doc, "The key is having one place where ideas go. Otherwise they disappear and you forget what you wanted to ask.")

    save_doc(doc, "05-RECOMMENDED-TOOLS-AND-SETUP.docx")


# ─── Document 3: Your First Week Framework ───

def create_first_week():
    doc = new_doc()

    add_heading(doc, "Your First Week: Step-by-Step Framework")
    add_para(doc, "Exactly what to do each day so you are set up and getting results within 7 days.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_para(doc, "Do not try to set up all 34 projects at once. That is overwhelming and unnecessary. This framework gets you running with your most important projects first, then you expand from there.", bold=True)
    doc.add_paragraph("")

    # Day 1
    add_heading(doc, "Day 1: Foundation", level=2)
    add_para(doc, "Goal: Get your Claude account ready and understand how Projects work.", bold=True)
    doc.add_paragraph("")
    add_checkbox(doc, "Upgrade to Claude Pro (or Max) at claude.ai")
    add_checkbox(doc, "Read 02-START-HERE in the onboarding folder")
    add_checkbox(doc, "Read 04-WHY-YOU-NEED-PRO-AND-MAX in the onboarding folder")
    add_checkbox(doc, "Install Wispr Flow or set up voice dictation on your device")
    add_checkbox(doc, "Create your first test project (use any template) just to get familiar with the interface")
    add_checkbox(doc, "Practice: Open the project, paste instructions, upload a file, and start a chat")
    doc.add_paragraph("")
    add_para(doc, "By the end of Day 1 you should be comfortable creating a project, pasting instructions, and chatting with Claude inside a project.")
    doc.add_paragraph("")

    # Day 2
    add_heading(doc, "Day 2: Voice and Brand Files", level=2)
    add_para(doc, "Goal: Create the knowledge files that every project needs.", bold=True)
    doc.add_paragraph("")
    add_checkbox(doc, "Record a 5-10 minute voice memo explaining your business (who you help, how, and why)")
    add_checkbox(doc, "Transcribe it and save as a PDF or text file")
    add_checkbox(doc, "Gather your 5-10 best pieces of content (captions, emails, posts)")
    add_checkbox(doc, "Save them in a single document called 'Content Samples'")
    add_checkbox(doc, "Create a one-page 'Business Overview' document (what you sell, who to, prices, website)")
    doc.add_paragraph("")
    add_para(doc, "These three files are your foundation. You will upload them to almost every project.")
    doc.add_paragraph("")

    # Day 3
    add_heading(doc, "Day 3: Your First Real Project", level=2)
    add_para(doc, "Goal: Set up the project that will save you the most time this week.", bold=True)
    doc.add_paragraph("")
    add_para(doc, "Pick ONE project based on what you need most right now:")
    add_bullet(doc, "Drowning in content creation? Start with Project 5 (Instagram Content Creator)")
    add_bullet(doc, "Need more leads? Start with Project 11 (DM & Outreach Sequences)")
    add_bullet(doc, "Spending hours on admin? Start with Project 28 (Personal Productivity & Planning)")
    add_bullet(doc, "Onboarding clients is messy? Start with Project 13 (Client Onboarding Assistant)")
    doc.add_paragraph("")
    add_checkbox(doc, "Open the project file from this pack")
    add_checkbox(doc, "Replace all [PLACEHOLDER] fields with your details")
    add_checkbox(doc, "Paste the instructions into your Claude Project")
    add_checkbox(doc, "Upload your voice file, content samples, and business overview")
    add_checkbox(doc, "Upload any project-specific knowledge files from the checklist")
    add_checkbox(doc, "Start a conversation and give it a real task")
    doc.add_paragraph("")
    add_para(doc, "Spend at least 30 minutes actually using this project. Give it real work. See what it produces.")
    doc.add_paragraph("")

    # Day 4
    add_heading(doc, "Day 4: Refine and Add a Second Project", level=2)
    add_para(doc, "Goal: Improve your first project and set up a second one.", bold=True)
    doc.add_paragraph("")
    add_checkbox(doc, "Review what your first project produced yesterday")
    add_checkbox(doc, "If the tone was off, add more voice samples or refine instructions")
    add_checkbox(doc, "If it missed context, upload additional knowledge files")
    add_checkbox(doc, "Give it another real task and compare the improvement")
    add_checkbox(doc, "Set up your second project following the same process")
    doc.add_paragraph("")
    add_para(doc, "This is where you start to see the power. Each project gets better the more context you give it.")
    doc.add_paragraph("")

    # Day 5
    add_heading(doc, "Day 5: Build Your Daily Workflow", level=2)
    add_para(doc, "Goal: Make Claude part of your actual workday, not just an experiment.", bold=True)
    doc.add_paragraph("")
    add_checkbox(doc, "Identify 3 tasks you do every week that Claude could handle")
    add_checkbox(doc, "Match each task to a project in this pack")
    add_checkbox(doc, "Set up a third project (or refine the first two)")
    add_checkbox(doc, "Create a 'Claude bookmark folder' in your browser with your active projects")
    add_checkbox(doc, "Try using voice (Wispr Flow or dictation) for at least 50% of your prompts today")
    doc.add_paragraph("")

    # Day 6-7
    add_heading(doc, "Days 6-7: Expand and Consolidate", level=2)
    add_para(doc, "Goal: Have 3-5 active projects and a clear routine.", bold=True)
    doc.add_paragraph("")
    add_checkbox(doc, "Set up 1-2 more projects that match your business priorities")
    add_checkbox(doc, "Review the Project Review Checklist to make sure your projects are properly configured")
    add_checkbox(doc, "Read the Cost Savings Breakdown to understand the full value of what you have built")
    add_checkbox(doc, "Write down your 'Claude routine': which projects you use, when, and for what")
    add_checkbox(doc, "Plan which projects you will add next week")
    doc.add_paragraph("")

    # After week 1
    add_heading(doc, "After Week 1: What Comes Next", level=2)
    add_para(doc, "You now have a working system. From here:")
    doc.add_paragraph("")
    add_bullet(doc, "Add 1-2 new projects per week as you identify new needs")
    add_bullet(doc, "Keep refining your knowledge files as your business evolves")
    add_bullet(doc, "Review and update project instructions quarterly")
    add_bullet(doc, "Share projects with team members or VAs if you have them")
    doc.add_paragraph("")

    # Month 1 rollout
    doc.add_page_break()
    add_heading(doc, "Month 1: Your Core 10 Projects", level=2)
    add_para(doc, "By the end of your first month, aim to have these 10 projects active. They are listed in the order you should set them up, based on what gives you the fastest return.", bold=True)
    doc.add_paragraph("")

    month1_projects = [
        ("Project 28: Personal Productivity & Planning", "Your daily command centre. No knowledge files needed to start. Set this up first for a quick win"),
        ("Project 5: Instagram Content Creator", "Stops you staring at a blank screen every time you need to post. Upload your best captions and voice memo"),
        ("Project 3: Email Marketing & Sequences", "Your email list is your most valuable asset. Get Claude writing sequences that sound like you"),
        ("Project 1: Content Strategy & Calendar", "Plan your entire month in one session instead of scrambling daily"),
        ("Project 9: Sales Page & Landing Page Copy", "Every offer needs a sales page. This writes high-converting copy in your voice"),
        ("Project 11: DM & Outreach Sequences", "Turn cold contacts into warm conversations without spending hours in the DMs"),
        ("Project 13: Client Onboarding Assistant", "First impressions matter. Make every new client feel like a VIP from day one"),
        ("Project 8: Content Repurposing Engine", "One piece of content becomes five. Stop creating from scratch every time"),
        ("Project 17: SOP & Process Documentation", "Document how you do things so you can delegate or automate later"),
        ("Project 30: Offer Design & Pricing Strategy", "Make sure your offers are structured and priced to actually make you money"),
    ]

    for i, (name, reason) in enumerate(month1_projects, 1):
        add_para(doc, f"{i}. {name}", bold=True)
        add_para(doc, reason)
        doc.add_paragraph("")

    add_para(doc, "Set up 2-3 per week. Do not rush all 10 at once. Each project needs proper customisation and testing to work well.", bold=True)
    doc.add_paragraph("")

    # Quarter 1 rollout
    doc.add_page_break()
    add_heading(doc, "Quarter 1: Full System Rollout (20+ Projects)", level=2)
    add_para(doc, "Once your core 10 are running, expand into these areas based on what your business needs most:")
    doc.add_paragraph("")

    add_para(doc, "Revenue & Growth (add these when you want to scale):", bold=True)
    add_bullet(doc, "Project 4: Lead Magnets & Opt-in Content")
    add_bullet(doc, "Project 10: Proposal & Pitch Writer")
    add_bullet(doc, "Project 12: Webinar & Workshop Planner")
    add_bullet(doc, "Project 22: Brand & Positioning Strategist")
    add_bullet(doc, "Project 24: Customer Research & Insights")
    doc.add_paragraph("")

    add_para(doc, "Content & Visibility (add these when you want to grow your audience):", bold=True)
    add_bullet(doc, "Project 2: Blog & Long-Form Content")
    add_bullet(doc, "Project 6: LinkedIn Thought Leadership")
    add_bullet(doc, "Project 7: YouTube Content & Scripts")
    add_bullet(doc, "Project 32: Podcast Strategy & Guest Appearances")
    add_bullet(doc, "Project 33: Community Management")
    doc.add_paragraph("")

    add_para(doc, "Operations & Admin (add these when you want to get your time back):", bold=True)
    add_bullet(doc, "Project 14: Coaching Session Prep & Notes")
    add_bullet(doc, "Project 15: Client Reporting & Results")
    add_bullet(doc, "Project 16: Course & Programme Content")
    add_bullet(doc, "Project 18: KPI & Business Reporting")
    add_bullet(doc, "Project 25: Email & Communication Manager")
    doc.add_paragraph("")

    add_para(doc, "Advanced (add these when your system is mature):", bold=True)
    add_bullet(doc, "Project 19: Team Communication & Management")
    add_bullet(doc, "Project 20: Finance & Invoicing Assistant")
    add_bullet(doc, "Project 21: Business Strategy & Planning")
    add_bullet(doc, "Project 23: Partnership & Collaboration")
    add_bullet(doc, "Project 26: Meeting & Event Coordinator")
    add_bullet(doc, "Project 27: Legal & Compliance Drafts")
    add_bullet(doc, "Project 29: Finance & Strategic Reporting")
    add_bullet(doc, "Project 31: Client Offboarding & Retention")
    add_bullet(doc, "Project 34: Tech Stack & Automation Setup")
    doc.add_paragraph("")

    add_para(doc, "By the end of Quarter 1 you should be running 15-20+ projects and saving 10-20 hours per week. That is not an exaggeration. That is what happens when you stop doing everything manually.", bold=True)

    save_doc(doc, "07-YOUR-FIRST-WEEK-FRAMEWORK.docx")


# ─── Document 4: How to Write Great Prompts ───

def create_prompting_guide():
    doc = new_doc()

    add_heading(doc, "How to Get the Best Results from Claude")
    add_para(doc, "Simple rules for writing prompts that actually get you what you want, first time.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_para(doc, "You do not need to learn 'prompt engineering'. You just need to talk to Claude like you would brief a team member. Here is how.", bold=True)
    doc.add_paragraph("")

    # Rule 1
    add_heading(doc, "Rule 1: Be Specific About What You Want", level=2)
    add_para(doc, "Bad prompt:", bold=True)
    add_para(doc, '"Write me an Instagram caption"')
    doc.add_paragraph("")
    add_para(doc, "Good prompt:", bold=True)
    add_para(doc, '"Write me an Instagram caption about how most business owners undercharge because they are scared of rejection. Target audience is business owners making 50-80K who know they should charge more but keep discounting. End with a soft CTA to DM me the word PRICING."')
    doc.add_paragraph("")
    add_para(doc, "The more context you give, the less back-and-forth you need. Think about how you would brief a copywriter. You would not just say 'write me something'. You would tell them the topic, the audience, the angle, and the goal.")
    doc.add_paragraph("")

    # Rule 2
    add_heading(doc, "Rule 2: Give Examples of What Good Looks Like", level=2)
    add_para(doc, "If you want Claude to match a specific style, show it an example:")
    doc.add_paragraph("")
    add_para(doc, '"Here is a caption I wrote last week that performed really well: [paste example]. Write me 3 more in this exact same style but on different topics."')
    doc.add_paragraph("")
    add_para(doc, "This is why uploading your best content as knowledge files makes such a difference. Claude can reference your examples without you pasting them every time.")
    doc.add_paragraph("")

    # Rule 3
    add_heading(doc, "Rule 3: Tell Claude What NOT to Do", level=2)
    add_para(doc, "Claude responds well to boundaries:")
    doc.add_paragraph("")
    add_bullet(doc, '"Do not use emojis"')
    add_bullet(doc, '"Do not start with a question"')
    add_bullet(doc, '"Avoid corporate jargon"')
    add_bullet(doc, '"Do not use bullet points, write in flowing paragraphs"')
    add_bullet(doc, '"Keep it under 150 words"')
    doc.add_paragraph("")
    add_para(doc, "The instructions in each project file already handle the big rules. But for individual prompts, adding a quick 'do not' can sharpen the output significantly.")
    doc.add_paragraph("")

    # Rule 4
    add_heading(doc, "Rule 4: Use Follow-Up Prompts to Refine", level=2)
    add_para(doc, "You do not need to get it perfect first time. Use follow-ups:")
    doc.add_paragraph("")
    add_bullet(doc, '"That is good but make the opening line more punchy"')
    add_bullet(doc, '"Rewrite this but more casual, like I am talking to a mate"')
    add_bullet(doc, '"Keep the structure but change the example to something about pricing"')
    add_bullet(doc, '"Shorter. Cut it in half"')
    add_bullet(doc, '"This sounds too AI. Make it sound more like my voice samples"')
    doc.add_paragraph("")
    add_para(doc, "Think of it as a conversation, not a one-shot request. The first output is a starting point. A couple of follow-ups usually gets you exactly what you need.")
    doc.add_paragraph("")

    # Rule 5
    add_heading(doc, "Rule 5: Use Voice Input (Seriously)", level=2)
    add_para(doc, "When you type, you tend to write prompts that are short and vague because typing is effort.")
    doc.add_paragraph("")
    add_para(doc, "When you speak, you naturally explain more. You give more context. You describe what you actually want. And that extra context is exactly what Claude needs to give you better output.")
    doc.add_paragraph("")
    add_para(doc, "Use Wispr Flow or your device's built-in voice typing for at least half your prompts. You will notice the difference immediately.", bold=True)
    doc.add_paragraph("")

    # Starter prompts
    add_heading(doc, "Starter Prompts to Try Right Now", level=2)
    add_para(doc, "Copy any of these into a project and see what happens:")
    doc.add_paragraph("")
    add_bullet(doc, '"Look at my voice samples and content examples. Write me 5 Instagram captions on [TOPIC] that sound exactly like me."')
    add_bullet(doc, '"I have a discovery call with a [TYPE OF COACH] tomorrow. Prepare me a list of questions to ask and objections they might have."')
    add_bullet(doc, '"Write me a 5-email welcome sequence for new subscribers. My lead magnet is [DESCRIBE]. The goal is to build trust and lead to a discovery call."')
    add_bullet(doc, '"Review my current offer and tell me where the gaps are. Here is what I sell: [DESCRIBE]"')
    add_bullet(doc, '"Create a weekly content plan for next week. My focus topics are [TOPICS] and I am active on Instagram and LinkedIn."')

    save_doc(doc, "08-HOW-TO-GET-BEST-RESULTS.docx")


# ─── Document 5: Troubleshooting & FAQs ───

def create_troubleshooting():
    doc = new_doc()

    add_heading(doc, "Troubleshooting and FAQs")
    add_para(doc, "Quick fixes for the most common issues people run into when setting up and using their Claude Projects.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Issue 1
    add_heading(doc, '"Claude does not sound like me"', level=2)
    add_para(doc, "This is the most common issue and the easiest to fix.")
    doc.add_paragraph("")
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Upload more voice samples. The more examples Claude has of how you actually write and speak, the better it matches your style")
    add_bullet(doc, "Record yourself talking about your business for 5-10 minutes, transcribe it, and upload it")
    add_bullet(doc, "Add a line to your instructions: 'Before writing anything, review my voice samples and match my exact tone and style'")
    add_bullet(doc, "When Claude produces something that sounds off, tell it specifically what is wrong: 'That sounds too formal. I would say it more like this...'")
    doc.add_paragraph("")

    # Issue 2
    add_heading(doc, '"I keep running out of messages"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Upgrade to Max plan if you are on Pro and consistently hitting limits")
    add_bullet(doc, "Write longer, more detailed prompts instead of lots of short back-and-forth messages")
    add_bullet(doc, "Batch your work: do all your content creation in one session, then all your operations work in another")
    add_bullet(doc, "Use the Sonnet model for simpler tasks (drafting, formatting) and save Opus for complex work (strategy, analysis)")
    doc.add_paragraph("")

    # Issue 3
    add_heading(doc, '"Claude forgot my instructions"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Make sure you are chatting inside the project, not in the general chat. Check the project name appears at the top of your conversation")
    add_bullet(doc, "Very long conversations can cause Claude to lose focus on earlier instructions. Start a fresh chat within the same project when this happens")
    add_bullet(doc, "Your custom instructions and knowledge files persist across all chats within a project, so starting fresh does not lose your setup")
    doc.add_paragraph("")

    # Issue 4
    add_heading(doc, '"The output is too generic"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Upload more specific knowledge files: real client examples, actual numbers, specific case studies")
    add_bullet(doc, "Be more specific in your prompts: include your niche, audience details, and the specific outcome you want")
    add_bullet(doc, "Tell Claude what makes your approach different from everyone else in your space")
    add_bullet(doc, "Replace [PLACEHOLDER] fields in your instructions with actual details, not vague descriptions")
    doc.add_paragraph("")

    # Issue 5
    add_heading(doc, '"I cannot find Projects in Claude"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Projects are only available on Pro and Max plans. Check your subscription at claude.ai under Settings")
    add_bullet(doc, "Projects appear in the left sidebar. If your sidebar is collapsed, click the menu icon to expand it")
    add_bullet(doc, "Make sure you are on claude.ai (the website), not the Claude mobile app. Projects are currently web-only for setup")
    doc.add_paragraph("")

    # Issue 6
    add_heading(doc, '"My file will not upload"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Check the file size. Knowledge files have a per-file limit and a total project limit of 200MB")
    add_bullet(doc, "Supported formats: PDF, TXT, DOCX, CSV, and common document formats")
    add_bullet(doc, "If your file is too large, split it into smaller files or extract the most important sections")
    add_bullet(doc, "Rename files clearly before uploading so Claude understands what each one contains")
    doc.add_paragraph("")

    # Issue 7
    add_heading(doc, '"I do not know which project to use"', level=2)
    add_para(doc, "Fix:", bold=True)
    add_bullet(doc, "Think about the task, not the category. What are you trying to produce or achieve right now?")
    add_bullet(doc, "Use the Project Review Checklist to see all 34 projects at a glance and what each one does")
    add_bullet(doc, "When in doubt, start with the project that matches the type of output you need: content, strategy, operations, or admin")
    add_bullet(doc, "It is fine to try a task in one project and move it to another if it fits better")
    doc.add_paragraph("")

    # General tips
    add_heading(doc, "General Tips", level=2)
    add_bullet(doc, "Start a new chat within a project when conversations get long. Your setup carries over automatically")
    add_bullet(doc, "Name your chats clearly so you can find past conversations")
    add_bullet(doc, "Keep knowledge files up to date. If your offers or prices change, update the files")
    add_bullet(doc, "Do not be afraid to experiment. You cannot break anything. Try different prompts and see what works best")

    save_doc(doc, "11-TROUBLESHOOTING-AND-FAQS.docx")


# ─── Document 6: Progress Tracker ───

def create_progress_tracker():
    doc = new_doc()

    add_heading(doc, "Your Setup Progress Tracker")
    add_para(doc, "Track your progress as you set up your Claude Projects system. Print this out or tick them off digitally.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Phase 1
    add_heading(doc, "Phase 1: Account and Tools Ready", level=2)
    add_checkbox(doc, "Upgraded to Claude Pro or Max")
    add_checkbox(doc, "Can access Claude Projects in the sidebar")
    add_checkbox(doc, "Installed Wispr Flow (or set up voice dictation)")
    add_checkbox(doc, "Pinned Claude as a browser tab")
    add_checkbox(doc, "Read the Start Here document")
    add_checkbox(doc, "Read the Why You Need Pro & Max document")
    doc.add_paragraph("")

    # Phase 2
    add_heading(doc, "Phase 2: Foundation Files Created", level=2)
    add_checkbox(doc, "Recorded and transcribed a voice memo about your business")
    add_checkbox(doc, "Compiled your 5-10 best content samples into one document")
    add_checkbox(doc, "Created a Business Overview document (offers, audience, prices, website)")
    add_checkbox(doc, "Saved all foundation files somewhere easy to find for uploading")
    doc.add_paragraph("")

    # Phase 3
    add_heading(doc, "Phase 3: First Projects Set Up", level=2)
    add_checkbox(doc, "Set up Project 1: ___________________________")
    add_checkbox(doc, "Set up Project 2: ___________________________")
    add_checkbox(doc, "Set up Project 3: ___________________________")
    add_checkbox(doc, "Tested each project with a real task")
    add_checkbox(doc, "Refined instructions or knowledge files based on output quality")
    doc.add_paragraph("")

    # Phase 4
    add_heading(doc, "Phase 4: Daily Workflow Established", level=2)
    add_checkbox(doc, "Identified 3+ weekly tasks that Claude handles for you")
    add_checkbox(doc, "Created browser bookmarks for your active projects")
    add_checkbox(doc, "Using voice input for at least some of your prompts")
    add_checkbox(doc, "Comfortable starting new chats and switching between projects")
    doc.add_paragraph("")

    # Phase 5
    add_heading(doc, "Phase 5: System Expansion", level=2)
    add_checkbox(doc, "Running 5+ active projects")
    add_checkbox(doc, "Reviewed the Cost Savings Breakdown")
    add_checkbox(doc, "Completed the Project Review Checklist for all active projects")
    add_checkbox(doc, "Planned which projects to add next")
    add_checkbox(doc, "Shared the system with team members or VA (if applicable)")
    doc.add_paragraph("")

    # Notes
    add_heading(doc, "Notes and Wins", level=2)
    add_para(doc, "Write down your biggest wins, time saved, or anything you want to remember as you build out your system:")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70)

    save_doc(doc, "10-SETUP-PROGRESS-TRACKER.docx")


# ─── Document 7: Your First 15 Minutes ───

def create_first_15_minutes():
    doc = new_doc()

    add_heading(doc, "Your First 15 Minutes: One Project, One Win")
    add_para(doc, "Set up your first Claude Project right now. No knowledge files needed. Just paste, personalise, and start chatting.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_para(doc, "Before you read anything else in this pack, do this. It takes 15 minutes and you will see exactly what this system can do for you.", bold=True)
    doc.add_paragraph("")

    add_heading(doc, "What You Are Setting Up", level=2)
    add_para(doc, "Project 28: Personal Productivity & Planning")
    add_para(doc, "This is your daily command centre. It helps you plan your week, prioritise tasks, set goals, and stay focused. It does not need any knowledge files to work. Just your basic business details typed into the instructions.")
    doc.add_paragraph("")

    add_heading(doc, "Step 1: Open Claude Projects", level=2)
    add_para(doc, "Go to claude.ai and log in. Click 'Projects' in the left sidebar. If you do not see Projects, you need to upgrade to Claude Pro first (see document 04 in this folder).")
    doc.add_paragraph("")

    add_heading(doc, "Step 2: Create a New Project", level=2)
    add_para(doc, "Click 'Create Project'. Set the name to:")
    add_para(doc, "Personal Productivity & Planning", bold=True)
    add_para(doc, "Add this description: 'Weekly planning, daily priorities, goal setting, and task management'")
    doc.add_paragraph("")

    add_heading(doc, "Step 3: Paste These Instructions", level=2)
    add_para(doc, "Copy everything in the box below. Before pasting, replace the 3 fields in [BRACKETS] with your own details. Then paste it into the 'Custom Instructions' box in your project.")
    doc.add_paragraph("")

    instructions = """You are my Personal Productivity Assistant and Planning Partner.

My business: [YOUR BUSINESS NAME - e.g. "Sarah's Fitness Coaching"]
What I do: [WHAT YOU DO - e.g. "I help busy professionals get fit with online coaching programmes"]
My biggest priorities right now: [YOUR TOP 3 PRIORITIES - e.g. "1. Launch my new group programme 2. Grow Instagram to 5K 3. Streamline client onboarding"]

Your role:
- Help me plan my week every Monday with clear priorities and time blocks
- When I brain dump tasks, organise them by priority and category
- Help me set realistic daily goals based on what actually moves the needle
- Challenge me if I am overloading my schedule or avoiding important tasks
- Keep me focused on revenue-generating activities, not just busy work
- Help me reflect on what worked and what did not at the end of each week

Rules:
- Be direct. Do not sugarcoat it if I am wasting time on low-value tasks
- Keep plans simple and realistic for a solo operator or small team
- Always ask what my top 3 priorities are before building any plan
- When I feel overwhelmed, help me ruthlessly cut or delegate
- Format plans as clear checklists I can work through, not essays"""

    for line in instructions.strip().split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")

    add_heading(doc, "Step 4: Start Chatting", level=2)
    add_para(doc, "Click 'Start Chat' and try one of these prompts:")
    doc.add_paragraph("")
    add_para(doc, "Prompt 1: Plan your week", bold=True)
    add_para(doc, '"Plan my week. I have 5 working days and about 6 hours of productive time each day. My top priorities are [list your 3 priorities]. What should each day look like?"')
    doc.add_paragraph("")
    add_para(doc, "Prompt 2: Brain dump", bold=True)
    add_para(doc, '"I have all of this swirling around my head: [list everything on your mind]. Organise this for me. What is urgent, what is important, and what can wait?"')
    doc.add_paragraph("")
    add_para(doc, "Prompt 3: Weekly review", bold=True)
    add_para(doc, '"Help me do a quick weekly review. This week I [describe what you did]. What went well, what did I avoid, and what should I focus on next week?"')
    doc.add_paragraph("")

    add_heading(doc, "That Is It. You Are Done.", level=2)
    add_para(doc, "You just set up your first Claude Project in under 15 minutes.", bold=True)
    doc.add_paragraph("")
    add_para(doc, "This is one of 34 projects in this pack. Every single one works the same way: paste the instructions, add your details, upload any knowledge files listed, and start chatting.")
    doc.add_paragraph("")
    add_para(doc, "Now open 02-START-HERE.docx to understand the full system and plan which projects to set up next.")
    doc.add_paragraph("")

    add_heading(doc, "Want to Watch Instead of Read?", level=2)
    add_para(doc, "A video walkthrough of this exact setup is available at: [VIDEO URL]")
    add_para(doc, "If the link above is not yet active, follow the steps on this page. They are identical to what the video shows.")

    save_doc(doc, "01-YOUR-FIRST-15-MINUTES.docx")


# ─── Document 8: Define Your Brand First ───

def create_brand_clarity():
    doc = new_doc()

    add_heading(doc, "Define Your Brand First")
    add_para(doc, "Fill this in once. Use it every time you set up a new project. No more guessing what to put in the placeholders.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_para(doc, "Every project template in this pack has placeholders in [SQUARE BRACKETS]. Instead of figuring out what to write each time, fill in this worksheet once and copy your answers into every project you set up.", bold=True)
    doc.add_paragraph("")

    add_heading(doc, "Your 7 Key Details", level=2)
    add_para(doc, "Fill in each field below. Write it how you would actually say it, not how you think it should sound.")
    doc.add_paragraph("")

    fields = [
        ("[YOUR BRAND NAME]", "Your business or personal brand name", "e.g. 'Sarah's Coaching Co' or 'The Fit Business' or just 'Sarah Mitchell'"),
        ("[YOUR NICHE]", "The specific group of people you help", "e.g. 'fitness business owners', 'wellness practitioners for women over 40', 'business consultants in the property space'"),
        ("[YOUR AUDIENCE]", "Describe your ideal client in one sentence", "e.g. 'business owners making 50-100K who want to scale without burning out', 'new personal trainers building their first online offer'"),
        ("[YOUR TONE]", "How you naturally communicate", "e.g. 'warm and conversational', 'direct and no-nonsense', 'energetic and motivating', 'calm and professional'"),
        ("[YOUR CTA]", "What you want people to do after reading your content", "e.g. 'Book a free discovery call at www.yoursite.com/call', 'DM me the word START', 'Download my free guide at www.yoursite.com/guide'"),
        ("[YOUR WEBSITE]", "Your main website URL", "e.g. 'www.yoursite.com'"),
        ("[YOUR OFFER]", "Your main product or service", "e.g. 'The Scale System (12-week group coaching programme)', '1:1 Business Mentoring', 'The Launch Accelerator Course'"),
    ]

    for placeholder, description, example in fields:
        add_para(doc, placeholder, bold=True)
        add_para(doc, description)
        add_para(doc, example, size=10, color=RGBColor(0x55, 0x55, 0x55))
        doc.add_paragraph("")
        add_para(doc, "Your answer: _______________________________________________")
        doc.add_paragraph("")

    doc.add_page_break()
    add_heading(doc, "Your Brand Clarity Cheat Sheet", level=2)
    add_para(doc, "Once you have filled in all 7 fields above, copy them here for quick reference:")
    doc.add_paragraph("")

    add_table(doc,
        ['Placeholder', 'Your Details'],
        [
            ['[YOUR BRAND NAME]', ''],
            ['[YOUR NICHE]', ''],
            ['[YOUR AUDIENCE]', ''],
            ['[YOUR TONE]', ''],
            ['[YOUR CTA]', ''],
            ['[YOUR WEBSITE]', ''],
            ['[YOUR OFFER]', ''],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Keep this page open or printed out whenever you are setting up a new project. Copy-paste from here instead of typing it fresh every time.", bold=True)
    doc.add_paragraph("")

    add_heading(doc, "Not Sure What to Write? Let Claude Help.", level=2)
    add_para(doc, "If you are struggling to define any of these, paste this prompt into any Claude conversation:")
    doc.add_paragraph("")

    prompt_text = "Help me define my brand positioning. I am a [describe what you do] and I help [describe who you help]. Ask me 10 questions about my business, who I serve, how I communicate, and what makes me different. Then give me a one-page summary I can use to fill in templates for my AI projects. Include: my brand name, niche, target audience description, communication tone, main CTA, website, and main offer."

    p = doc.add_paragraph(prompt_text)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")
    add_para(doc, "Claude will ask you follow-up questions and then generate a complete brand summary you can use to fill in this worksheet.")

    save_doc(doc, "03-DEFINE-YOUR-BRAND-FIRST.docx")


# ─── Document 9: Knowledge File Starter Kit ───

def create_knowledge_starter_kit():
    doc = new_doc()

    add_heading(doc, "Knowledge File Starter Kit")
    add_para(doc, "You do not need to create 50 documents from scratch. You already have most of what you need. This guide shows you where to find it and how to prepare it.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_heading(doc, "The 3 Files That Power Everything", level=2)
    add_para(doc, "If you only upload 3 files to your projects, upload these. They cover 80% of what Claude needs to produce great output.", bold=True)
    doc.add_paragraph("")

    add_para(doc, "File 1: Voice Memo Transcript", bold=True)
    add_bullet(doc, "What it is: A transcription of you talking naturally about your business for 5-10 minutes")
    add_bullet(doc, "How to create it: Open the voice recorder on your phone. Talk about what your business does, who you help, what makes you different, and why you started. Do not script it. Just talk like you are explaining it to a friend")
    add_bullet(doc, "How to transcribe: Use Wispr Flow, your phone's built-in dictation, or paste the audio into a free transcription tool like Otter.ai")
    add_bullet(doc, "Save as: Voice-Memo-Transcript.txt or .pdf")
    add_bullet(doc, "Upload to: Every project that creates written content (Instagram, email, blog, LinkedIn, etc.)")
    doc.add_paragraph("")

    add_para(doc, "File 2: Content Samples Document", bold=True)
    add_bullet(doc, "What it is: A collection of your 5-10 best pieces of content in one document")
    add_bullet(doc, "How to create it: Open a Google Doc or Word doc. Copy and paste your best Instagram captions, emails, LinkedIn posts, or any written content you are proud of")
    add_bullet(doc, "What to include: Content that got good engagement, content that sounds most like you, content that clients have commented on")
    add_bullet(doc, "Save as: Content-Samples.docx or .pdf")
    add_bullet(doc, "Upload to: Every writing-based project")
    doc.add_paragraph("")

    add_para(doc, "File 3: Business Overview", bold=True)
    add_bullet(doc, "What it is: A simple one-page summary of your business")
    add_bullet(doc, "What to include: Your business name, what you sell (with prices), who you sell it to, your website, your main CTA, and what makes you different")
    add_bullet(doc, "How to create it: Open a document and write it in bullet points. It does not need to be polished")
    add_bullet(doc, "Save as: Business-Overview.docx or .pdf")
    add_bullet(doc, "Upload to: Every single project")
    doc.add_paragraph("")

    add_para(doc, "These 3 files will get you 80% of the way. You can always add more later.", bold=True, size=13)
    doc.add_paragraph("")

    doc.add_page_break()
    add_heading(doc, "You Already Have These (You Just Need to Find Them)", level=2)
    add_para(doc, "Most of the knowledge files in this pack are things you already have somewhere. Here is where to look:")
    doc.add_paragraph("")

    add_table(doc,
        ['Knowledge File Needed', 'Where You Probably Already Have It'],
        [
            ['Client testimonials', 'Google reviews, DMs, thank-you emails, Facebook group posts'],
            ['Sales page copy', 'Your current website (copy-paste the text)'],
            ['Email sequences', 'Export from Mailchimp, ConvertKit, ActiveCampaign, etc.'],
            ['Past proposals', 'Search your sent emails for "proposal" or "quote"'],
            ['SOPs or processes', 'That Notion page, Google Doc, or Loom library you already have'],
            ['Brand guidelines', 'Any document from your designer, Canva brand kit, or website brief'],
            ['Case studies', 'Client results you have shared on social media or in sales conversations'],
            ['Competitor analysis', 'Screenshots, bookmarks, or notes you have saved about competitors'],
            ['Pricing structure', 'Your website, Stripe dashboard, or that spreadsheet you track revenue in'],
            ['Client avatar', 'Any marketing course worksheet you have filled in, or your own description of your ideal client'],
        ]
    )
    doc.add_paragraph("")

    add_para(doc, "You do not need to create formal documents for all of these. A quick copy-paste from an email, website, or existing document works perfectly. Claude does not care about formatting. It cares about content.")
    doc.add_paragraph("")

    add_heading(doc, "Let Claude Create Them For You", level=2)
    add_para(doc, "Every project template in this pack has a section at the bottom called 'Generate Your Knowledge Files'. These are ready-to-use prompts that will create the files for you.")
    doc.add_paragraph("")
    add_para(doc, "Here is how it works:", bold=True)
    add_bullet(doc, "Open any Claude conversation (does not need to be inside a project)")
    add_bullet(doc, "Paste the generation prompt from the project template")
    add_bullet(doc, "Claude will ask you follow-up questions about your business")
    add_bullet(doc, "Answer the questions naturally (use voice input for best results)")
    add_bullet(doc, "Claude produces a complete document")
    add_bullet(doc, "Copy the output, save it as a .docx or .pdf, and upload it to the project")
    doc.add_paragraph("")
    add_para(doc, "Most knowledge files can be generated in 5-10 minutes each using this method.")
    doc.add_paragraph("")

    add_heading(doc, "Knowledge File Upload Priority", level=2)
    add_para(doc, "You do not need all the files listed in a project's checklist to get started. Here is the priority:")
    doc.add_paragraph("")

    add_para(doc, "Start with (essential):", bold=True)
    add_bullet(doc, "Your 3 foundation files (voice memo, content samples, business overview)")
    add_bullet(doc, "Any files you already have ready to go")
    doc.add_paragraph("")

    add_para(doc, "Add next (improves quality significantly):", bold=True)
    add_bullet(doc, "Project-specific files from the checklist (e.g. offer details for a sales page project)")
    add_bullet(doc, "Client testimonials or case studies")
    add_bullet(doc, "Any files generated using the built-in prompts")
    doc.add_paragraph("")

    add_para(doc, "Add over time (polishes the output):", bold=True)
    add_bullet(doc, "Competitor analysis documents")
    add_bullet(doc, "Historical data and performance reports")
    add_bullet(doc, "Detailed SOPs and workflow documentation")
    doc.add_paragraph("")

    add_para(doc, "The more context you give Claude, the better the output. But you can start with the basics and add more as you go. Do not let the checklist stop you from setting up a project today.", bold=True)

    save_doc(doc, "06-KNOWLEDGE-FILE-STARTER-KIT.docx")


# ─── Document 10: All 34 Projects at a Glance ───

def create_master_overview():
    doc = new_doc()

    add_heading(doc, "All 34 Projects at a Glance")
    add_para(doc, "Every project in this pack with a one-line description, ROI rating, and setup difficulty. Use this to decide what to set up next.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    # Top 5 callouts
    add_heading(doc, "Top 5 Quick Wins (Easiest Setup, Immediate Value)", level=2)
    add_bullet(doc, "Project 28: Personal Productivity & Planning \u2014 No knowledge files needed. Set up in 15 minutes")
    add_bullet(doc, "Project 5: Instagram Content Creator \u2014 Upload your best captions and go")
    add_bullet(doc, "Project 25: Email & Communication Manager \u2014 Templates for every email you send repeatedly")
    add_bullet(doc, "Project 3: Email Marketing & Sequences \u2014 Stop writing the same welcome emails from scratch")
    add_bullet(doc, "Project 8: Content Repurposing Engine \u2014 One piece of content becomes five")
    doc.add_paragraph("")

    add_heading(doc, "Top 5 Revenue Drivers (Directly Generate Income)", level=2)
    add_bullet(doc, "Project 9: Sales Page & Landing Page Copy \u2014 High-converting sales pages in your voice")
    add_bullet(doc, "Project 30: Offer Design & Pricing Strategy \u2014 Structure offers that actually sell")
    add_bullet(doc, "Project 11: DM & Outreach Sequences \u2014 Turn cold contacts into warm leads")
    add_bullet(doc, "Project 4: Lead Magnets & Opt-in Content \u2014 Build your list with valuable freebies")
    add_bullet(doc, "Project 10: Proposal & Pitch Writer \u2014 Win more clients with better proposals")
    doc.add_paragraph("")

    add_heading(doc, "Top 5 Time Savers (Biggest Hours Reclaimed)", level=2)
    add_bullet(doc, "Project 17: SOP & Process Documentation \u2014 Document everything so you can delegate")
    add_bullet(doc, "Project 1: Content Strategy & Calendar \u2014 Plan your whole month in one sitting")
    add_bullet(doc, "Project 13: Client Onboarding Assistant \u2014 Automate the welcome experience")
    add_bullet(doc, "Project 34: Tech Stack & Automation Setup \u2014 Streamline your tools and workflows")
    add_bullet(doc, "Project 14: Coaching Session Prep & Notes \u2014 Stop spending 30 minutes prepping each session")
    doc.add_paragraph("")

    # Full table
    doc.add_page_break()
    add_heading(doc, "Complete Project Directory", level=2)
    doc.add_paragraph("")

    # Data for all 34 projects
    projects_data = [
        # Content Marketing
        ("Content Marketing", [
            (1, "Content Strategy & Calendar", "Plan monthly themes and weekly content calendars", "High", "Easy"),
            (2, "Blog & Long-Form Content", "SEO articles, guides, and thought leadership pieces", "Medium", "Medium"),
            (3, "Email Marketing & Sequences", "Welcome sequences, nurture campaigns, launch emails", "High", "Easy"),
            (4, "Lead Magnets & Opt-in Content", "Checklists, guides, and downloadable resources", "High", "Medium"),
            (33, "Community Management", "Group rules, welcome sequences, engagement content", "Medium", "Medium"),
        ]),
        ("Social Media", [
            (5, "Instagram Content Creator", "Captions, carousels, reel scripts, stories", "High", "Easy"),
            (6, "LinkedIn Thought Leadership", "Posts, articles, connection messages", "Medium", "Medium"),
            (7, "YouTube Content & Scripts", "Full scripts with timestamps, titles, descriptions", "Medium", "Advanced"),
            (8, "Content Repurposing Engine", "Turn one piece of content into five platforms", "High", "Easy"),
            (32, "Podcast Strategy & Guest Appearances", "Topic lists, speaker bios, podcast pitches", "Medium", "Medium"),
        ]),
        ("Sales & Lead Generation", [
            (9, "Sales Page & Landing Page Copy", "High-converting sales pages and landing pages", "High", "Medium"),
            (10, "Proposal & Pitch Writer", "Client proposals, pitch decks, follow-ups", "High", "Medium"),
            (11, "DM & Outreach Sequences", "Personalised outreach and follow-up scripts", "High", "Easy"),
            (12, "Webinar & Workshop Planner", "Webinar structure, scripts, follow-up sequences", "Medium", "Advanced"),
        ]),
        ("Client Fulfilment", [
            (13, "Client Onboarding Assistant", "Welcome packets, onboarding emails, intake forms", "High", "Easy"),
            (14, "Coaching Session Prep & Notes", "Session agendas, questions, follow-up summaries", "High", "Medium"),
            (15, "Client Reporting & Results", "Monthly reports, progress summaries, ROI breakdowns", "Medium", "Medium"),
            (16, "Course & Programme Content", "Module outlines, lesson scripts, workbook pages", "Medium", "Advanced"),
            (31, "Client Offboarding & Retention", "Offboarding sequences, referral systems, alumni offers", "Medium", "Medium"),
        ]),
        ("Operations", [
            (17, "SOP & Process Documentation", "Step-by-step procedures and training guides", "High", "Easy"),
            (18, "KPI & Business Reporting", "Metrics analysis, performance reports, trend spotting", "Medium", "Medium"),
            (19, "Team Communication & Management", "Team updates, meeting agendas, feedback reviews", "Low", "Medium"),
            (20, "Finance & Invoicing Assistant", "Pricing structures, invoice templates, budgets", "Low", "Easy"),
            (29, "Finance & Strategic Reporting", "Revenue breakdowns, forecasting, CFO-level analysis", "Medium", "Advanced"),
            (34, "Tech Stack & Automation Setup", "Tool recommendations, automation workflows", "High", "Medium"),
        ]),
        ("Strategy & Growth", [
            (21, "Business Strategy & Planning", "Quarterly plans, goal setting, business vision", "Medium", "Medium"),
            (22, "Brand & Positioning Strategist", "Brand identity, positioning, differentiation", "High", "Medium"),
            (23, "Partnership & Collaboration", "JV partners, cross-promotion, collaboration planning", "Low", "Medium"),
            (24, "Customer Research & Insights", "Market research, customer insights, competitive analysis", "Medium", "Medium"),
            (30, "Offer Design & Pricing Strategy", "Offer architecture, pricing, offer waterfall", "High", "Medium"),
        ]),
        ("Admin & Organisation", [
            (25, "Email & Communication Manager", "Email templates, FAQs, communication standards", "Medium", "Easy"),
            (26, "Meeting & Event Coordinator", "Agendas, event checklists, follow-up templates", "Low", "Easy"),
            (27, "Legal & Compliance Drafts", "Terms, privacy policies, disclaimers, contracts", "Low", "Medium"),
            (28, "Personal Productivity & Planning", "Weekly planning, priorities, goal setting", "High", "Easy"),
        ]),
    ]

    for category, projects in projects_data:
        add_heading(doc, category, level=3)
        add_table(doc,
            ['#', 'Project Name', 'What It Does', 'ROI', 'Setup'],
            [[str(num), name, desc, roi, difficulty] for num, name, desc, roi, difficulty in projects]
        )
        doc.add_paragraph("")

    doc.add_paragraph("")
    add_para(doc, "ROI ratings: High = directly saves time or generates revenue. Medium = improves quality and consistency. Low = useful but not urgent.", size=10, color=RGBColor(0x55, 0x55, 0x55))
    add_para(doc, "Setup difficulty: Easy = minimal knowledge files, quick to set up. Medium = needs some knowledge files and customisation. Advanced = needs detailed knowledge files and thorough testing.", size=10, color=RGBColor(0x55, 0x55, 0x55))

    save_doc(doc, "09-ALL-34-PROJECTS-AT-A-GLANCE.docx")


# ─── Document 11: Project Review Checklist ───

def create_project_review_checklist():
    doc = new_doc()

    add_heading(doc, "Project Review Checklist")
    add_para(doc, "Use this checklist to track which projects you have set up and whether each one is properly configured.", size=13, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")

    add_para(doc, "Tick off each step as you complete it. A project is 'active' once all boxes are ticked for that project.", bold=True)
    doc.add_paragraph("")

    all_projects = [
        (1, "Content Strategy & Calendar", "01-Content-Marketing"),
        (2, "Blog & Long-Form Content", "01-Content-Marketing"),
        (3, "Email Marketing & Sequences", "01-Content-Marketing"),
        (4, "Lead Magnets & Opt-in Content", "01-Content-Marketing"),
        (5, "Instagram Content Creator", "02-Social-Media"),
        (6, "LinkedIn Thought Leadership", "02-Social-Media"),
        (7, "YouTube Content & Scripts", "02-Social-Media"),
        (8, "Content Repurposing Engine", "02-Social-Media"),
        (9, "Sales Page & Landing Page Copy", "03-Sales-Lead-Gen"),
        (10, "Proposal & Pitch Writer", "03-Sales-Lead-Gen"),
        (11, "DM & Outreach Sequences", "03-Sales-Lead-Gen"),
        (12, "Webinar & Workshop Planner", "03-Sales-Lead-Gen"),
        (13, "Client Onboarding Assistant", "04-Client-Fulfilment"),
        (14, "Coaching Session Prep & Notes", "04-Client-Fulfilment"),
        (15, "Client Reporting & Results", "04-Client-Fulfilment"),
        (16, "Course & Programme Content", "04-Client-Fulfilment"),
        (17, "SOP & Process Documentation", "05-Operations"),
        (18, "KPI & Business Reporting", "05-Operations"),
        (19, "Team Communication & Management", "05-Operations"),
        (20, "Finance & Invoicing Assistant", "05-Operations"),
        (21, "Business Strategy & Planning", "06-Strategy-Growth"),
        (22, "Brand & Positioning Strategist", "06-Strategy-Growth"),
        (23, "Partnership & Collaboration", "06-Strategy-Growth"),
        (24, "Customer Research & Insights", "06-Strategy-Growth"),
        (25, "Email & Communication Manager", "07-Admin-Organisation"),
        (26, "Meeting & Event Coordinator", "07-Admin-Organisation"),
        (27, "Legal & Compliance Drafts", "07-Admin-Organisation"),
        (28, "Personal Productivity & Planning", "07-Admin-Organisation"),
        (29, "Finance & Strategic Reporting", "05-Operations"),
        (30, "Offer Design & Pricing Strategy", "06-Strategy-Growth"),
        (31, "Client Offboarding & Retention", "04-Client-Fulfilment"),
        (32, "Podcast Strategy & Guest Appearances", "02-Social-Media"),
        (33, "Community Management", "01-Content-Marketing"),
        (34, "Tech Stack & Automation Setup", "05-Operations"),
    ]

    for num, title, _ in all_projects:
        add_para(doc, f"Project {num}: {title}", bold=True)
        add_checkbox(doc, "Custom instructions pasted and all [PLACEHOLDERS] replaced with my details")
        add_checkbox(doc, "Voice and brand files uploaded (if this is a writing project)")
        add_checkbox(doc, "Project-specific knowledge files uploaded")
        add_checkbox(doc, "Tested with a real task and reviewed the output")
        add_checkbox(doc, "Instructions refined based on output quality")
        doc.add_paragraph("")

    doc.add_page_break()
    add_heading(doc, "Summary", level=2)
    add_para(doc, "Projects Active: _______ / 34")
    doc.add_paragraph("")
    add_para(doc, "Date of last review: _______________")
    doc.add_paragraph("")
    add_para(doc, "Review your active projects quarterly to make sure instructions and knowledge files are still up to date. Use the Claude Code Guides folder (08) for bulk updates when things change.", bold=True)

    save_doc(doc, "12-PROJECT-REVIEW-CHECKLIST.docx")


# ─── Generate everything ───

def main():
    print("Generating Onboarding Documents...\n")

    os.makedirs(ONBOARDING_DIR, exist_ok=True)

    create_first_15_minutes()          # 01
    create_brand_clarity()             # 03
    create_plan_guide()                # 04
    create_tools_guide()               # 05
    create_knowledge_starter_kit()     # 06
    create_first_week()                # 07
    create_prompting_guide()           # 08
    create_master_overview()           # 09
    create_progress_tracker()          # 10
    create_troubleshooting()           # 11
    create_project_review_checklist()  # 12

    print(f"\nDone! Onboarding documents created in: {ONBOARDING_DIR}")


if __name__ == "__main__":
    main()
